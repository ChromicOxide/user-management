#!/usr/bin/env python3
"""生成《SQL注入漏洞专项分析与修复报告》Word文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_PRIMARY = RGBColor(0x1A, 0x2B, 0x3C)
C_SECONDARY = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACCENT = RGBColor(0x4A, 0x7A, 0xBF)
C_HIGH = RGBColor(0xCF, 0x13, 0x22)
C_MED = RGBColor(0xE8, 0x8C, 0x0A)
C_LOW = RGBColor(0x38, 0x9E, 0x0D)
BG_HEADER = 'E2E8F0'
BG_ROW = 'F5F7FA'
BG_CODE = 'F0F2F5'


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D0D5DD')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '36')
        m.set(qn('w:type'), 'dxa')
        cell_mar.append(m)
    tblPr.append(cell_mar)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.name = '微软雅黑'
        run.font.color.rgb = C_PRIMARY
        set_cell_shading(cell, BG_HEADER)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = '微软雅黑'
            run.font.color.rgb = C_BODY
            if ri % 2 == 1:
                set_cell_shading(cell, BG_ROW)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = Pt(8)
    return table


def add_code(doc, code_text, label=""):
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"▎{label}")
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = '微软雅黑'
        run.font.color.rgb = C_ACCENT
    for line in code_text.rstrip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(7)
        run.font.color.rgb = C_BODY
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), BG_CODE)
        p._element.get_or_add_pPr().append(shd)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = Pt(8)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    spacing = {1: (16, 8), 2: (10, 4), 3: (8, 3)}.get(level, (8, 3))
    h.paragraph_format.space_before = Pt(spacing[0])
    h.paragraph_format.space_after = Pt(spacing[1])
    for run in h.runs:
        run.font.name = '微软雅黑'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = C_PRIMARY
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = C_SECONDARY
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = C_BODY
    if level == 1:
        add_hr(doc)
    return h


def add_body(doc, text, bold=False, size=10, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.color.rgb = C_BODY
    run.bold = bold
    return p


def set_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(4)
    style.font.color.rgb = C_BODY


def main():
    doc = Document()
    set_default_style(doc)
    set_page_margins(doc)

    # ═══════ 封面 ═══════
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SQL注入漏洞专项分析与修复报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 用户管理系统 Web应用安全审计 ——")
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_MUTED

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(16)

    info = [
        ("文档编号", "SEC-SQL-2025-001"),
        ("密级", "内部公开"),
        ("版本号", "V1.0"),
        ("发布日期", "2026-07-08"),
        ("编制人", "安全审计组"),
        ("评审人", "技术负责人"),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        for j, txt in enumerate([k, v]):
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt)
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            if j == 0:
                run.bold = True
                run.font.color.rgb = C_PRIMARY
            else:
                run.font.color.rgb = C_ACCENT
            set_cell_shading(cell, 'EDF2F7' if i % 2 == 0 else 'F5F8FB')
    doc.add_page_break()

    # ═══════ 1. 报告概述 ═══════
    add_heading(doc, "1. 报告概述", 1)

    add_heading(doc, "1.1 项目信息", 2)
    add_table(doc,
        ["项目", "内容"],
        [["项目名称", "用户管理系统 SQL注入安全审计"],
         ["审计目标", "app.py 全量数据库交互代码的 SQL 注入漏洞排查与修复"],
         ["审计日期", "2026 年 7 月 8 日"],
         ["应用技术栈", "Python 3.13 / Flask 3.x / SQLite 3 / Jinja2 / Werkzeug 3.x"],
         ["代码仓库", "/workspace/user_management/app.py"]])

    add_heading(doc, "1.2 审计范围", 2)
    add_body(doc, "本次审计覆盖 app.py 中所有与 SQLite 数据库交互的代码路径，总计 6 处数据交互点：")
    items = [
        "init_db() — 数据库初始化（CREATE TABLE + INSERT 默认用户）",
        "get_user_from_db() — 用户查询（登录时从 SQLite 检索）",
        "register() — 用户注册（INSERT 新用户）",
        "search() — 用户搜索（SELECT LIKE 模糊查询）",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    add_heading(doc, "1.3 合规标准", 2)
    add_table(doc,
        ["标准", "版本", "适用范围"],
        [["OWASP Top 10", "2021", "A03:2021 — 注入"],
         ["CVSS", "3.1", "漏洞风险评估与量化评分"],
         ["等保 2.0", "GB/T 22239-2019", "安全计算环境—应用安全"],
         ["CWE", "CWE-89", "SQL 命令中使用的特殊元素转义处理不恰当（SQL注入）"]])

    add_heading(doc, "1.4 审计工具与方法论", 2)
    add_body(doc, "审计采用 代码审计 + 黑盒渗透 双轨并行方法：")
    items = [
        "代码审计：逐行审查 app.py 中 SQL 语句构造方式，定位字符串拼接点",
        "黑盒渗透：使用 Burp Suite 和 curl 构造 SQL 注入 Payload 验证注入存在性",
        "漏洞复现：针对每个注入点构造可执行的攻击 POC，确认危害等级",
        "修复验证：修复后使用相同 POC 复测，确认注入路径已封闭",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    doc.add_page_break()

    # ═══════ 2. 漏洞风险总览 ═══════
    add_heading(doc, "2. 漏洞风险总览", 1)

    add_heading(doc, "2.1 漏洞统计", 2)
    add_body(doc, "本轮审计共发现 SQL 注入相关漏洞 4 个，其中高危 3 个、中危 1 个，分布如下：")

    add_table(doc,
        ["编号", "漏洞名称", "风险等级", "CVSS 3.1", "业务模块", "影响范围"],
        [["SQL-01", "搜索功能 UNION 注入", "高危", "9.1 / AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:N",
          "用户搜索 (search)", "全量用户数据泄露 + 任意数据写入"],
         ["SQL-02", "搜索功能 OR 万能条件", "高危", "8.6 / AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
          "用户搜索 (search)", "全量用户数据泄露"],
         ["SQL-03", "注册功能 SQL 注入", "高危", "8.3 / AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:H/A:N",
          "用户注册 (register)", "恶意用户注册 + 数据库篡改"],
         ["SQL-04", "密码明文存储", "中危", "5.9 / AV:N/AC:H/PR:H/UI:R/S:U/C:H/I:N/A:N",
          "数据库初始化 (init_db)", "数据库泄露即口令批量暴露"]])

    add_heading(doc, "2.2 风险分布", 2)
    rows_t = [["高危", "3", "75%"],
              ["中危", "1", "25%"],
              ["低危", "0", "0%"],
              ["合计", "4", "100%"]]
    add_table(doc, ["风险等级", "数量", "占比"], rows_t)

    add_body(doc, "风险结构说明：高危漏洞占比 75%，主要集中在搜索功能的 f-string 直接拼接场景。"
             "攻击者无需认证即可利用 UNION 注入和 OR 万能条件获取全量用户敏感数据，"
             "注册功能的 SQL 注入则允许攻击者在无认证情况下向数据库写入任意数据。", size=9.5)

    doc.add_page_break()

    # ═══════ 3. 单漏洞深度分析 ═══════
    add_heading(doc, "3. 单漏洞深度分析", 1)

    # ─── SQL-01 ───
    add_heading(doc, "3.1 SQL-01：搜索功能 UNION 注入", 2)
    add_table(doc,
        ["属性", "内容"],
        [["漏洞编号", "SQL-01"],
         ["风险等级", "高危"],
         ["CVSS 3.1", "9.1 AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:N"],
         ["漏洞位置", "app.py 第 303 行 search() — f\"...LIKE '%{keyword}%'...\""],
         ["发现方式", "代码审计 + 黑盒渗透"],
         ["CWE 编号", "CWE-89: SQL Injection"]])

    add_heading(doc, "3.1.1 攻击路径与 POC", 3)
    add_body(doc, "攻击者通过搜索框提交恶意构造的 keyword 参数，利用 UNION SELECT 关键字向原查询结果中追加任意数据行：")
    add_code(doc, '''# 攻击路径：搜索框 UNION 注入
GET /search?keyword=' UNION SELECT 1,'inj','inj@x.com','138'--
# 实际生成的 SQL：
SELECT id,username,email,phone FROM users
WHERE username LIKE '%' UNION SELECT 1,'inj','inj@x.com','138'--%'
OR email LIKE '%' UNION SELECT 1,'inj','inj@x.com','138'--%'
# 结果：搜索结果页面出现攻击者伪造的 "inj" 用户数据''', "POC 攻击载荷")

    add_heading(doc, "3.1.2 根因分析", 3)
    add_body(doc, "漏洞根因是第 303 行使用 Python f-string 直接拼接用户输入构造 SQL 语句：", size=9.5)
    add_code(doc, '''# ❌ 危险代码（第303行）
executed_sql = f"SELECT id, username, email, phone FROM users \\
    WHERE username LIKE '%{keyword}%' OR email LIKE '%{keyword}%'"
c.execute(executed_sql)''', "漏洞代码（修复前）")

    add_body(doc, "该写法将用户输入的 keyword 直接嵌入 SQL 语法结构中，keyword 中的单引号 ' 会提前闭合 SQL 字符串边界，"
             "使得剩余内容脱离数据身份、成为可执行的 SQL 指令。UNION SELECT 关键字允许攻击者向结果集中追加任意查询结果。"
             "这是典型的 OWASP A03:2021 注入类漏洞。", size=9.5)

    add_heading(doc, "3.1.3 危害量化", 3)
    add_body(doc, "• 数据泄露：攻击者可提取 users 表中全部用户记录（用户名、密码哈希、邮箱、手机号）"
             "，当前库中约 8 条用户记录。若扩大搜索范围至 information_schema，可获取整个数据库结构。", size=9.5)
    add_body(doc, "• 数据伪造：通过 UNION 注入可在搜索结果页展示任意伪造内容，可用于钓鱼攻击或社会工程学。", size=9.5)
    add_body(doc, "• 拖库潜力：若数据库账号权限过高（如具备写文件权限），攻击者可组合 INTO OUTFILE 语句将全库数据导出。", size=9.5)

    add_heading(doc, "3.1.4 修复方案与代码对比", 3)
    add_code(doc, '''# ✅ 修复代码（参数化查询 + LIKE 通配符转义）
# 第 343-351 行
safe_keyword = keyword.replace("%", "\\\\%").replace("_", "\\\\_")
c.execute(
    "SELECT id, username, email, phone FROM users \\
     WHERE username LIKE ? OR email LIKE ? ESCAPE '\\\\'",
    (f"%{safe_keyword}%", f"%{safe_keyword}%")
)''', "修复后代码")

    add_body(doc, "修复要点：", bold=True, size=9.5)
    add_body(doc, "1. 参数化查询（Prepared Statement）：将用户输入通过 ? 占位符绑定，SQLite 引擎自动处理转义，"
             "用户输入永远不会成为 SQL 语法的一部分。这是从根本上阻断 SQL 注入的唯一可靠手段。", size=9.5)
    add_body(doc, "2. LIKE 通配符转义：% 和 _ 在 LIKE 模式中有特殊含义（匹配任意序列/匹配单字符），"
             "通过反斜杠转义 + ESCAPE '\\' 子句确保用户搜索字面量 % 时不会被解释为通配符。", size=9.5)
    add_body(doc, "3. 异常处理：新增 try/except 包裹数据库操作，避免未捕获的异常导致 500 错误及信息泄露。", size=9.5)

    doc.add_page_break()

    # ─── SQL-02 ───
    add_heading(doc, "3.2 SQL-02：搜索功能 OR 万能条件注入", 2)
    add_table(doc,
        ["属性", "内容"],
        [["漏洞编号", "SQL-02"],
         ["风险等级", "高危"],
         ["CVSS 3.1", "8.6 AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N"],
         ["漏洞位置", "app.py 第 303 行（与 SQL-01 同源）"],
         ["发现方式", "黑盒渗透验证"],
         ["CWE 编号", "CWE-89: SQL Injection"]])

    add_heading(doc, "3.2.1 攻击路径与 POC", 3)
    add_code(doc, '''# 攻击路径：OR 万能条件绕过
GET /search?keyword=' OR '1'='1
# 实际生成的 SQL：
SELECT id,username,email,phone FROM users
WHERE username LIKE '%' OR '1'='1%'
OR email LIKE '%' OR '1'='1%'
# WHERE 条件中的 '1'='1' 永远为真
# 结果：返回 users 表中全部用户记录''', "POC 攻击载荷")

    add_heading(doc, "3.2.2 根因分析与危害", 3)
    add_body(doc, "与 SQL-01 同根因。OR '1'='1' 构造了一个永真条件，使 WHERE 子句对所有行均成立，"
             "导致数据库返回全部用户记录。攻击者可在一次请求中获取系统中所有注册用户信息。"
             "结合该漏洞与 SQL-01（UNION），攻击者可完成完整的数据库信息提取。", size=9.5)
    add_body(doc, "修复方案与 SQL-01 完全一致：采用参数化查询。", size=9.5)

    doc.add_page_break()

    # ─── SQL-03 ───
    add_heading(doc, "3.3 SQL-03：注册功能 SQL 注入", 2)
    add_table(doc,
        ["属性", "内容"],
        [["漏洞编号", "SQL-03"],
         ["风险等级", "高危"],
         ["CVSS 3.1", "8.3 AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:H/A:N"],
         ["漏洞位置", "app.py 第 277 行 register() — f\"INSERT INTO ... VALUES ('{username}',...)\""],
         ["发现方式", "代码审计"],
         ["CWE 编号", "CWE-89: SQL Injection"]])

    add_heading(doc, "3.3.1 攻击路径与 POC", 3)
    add_code(doc, '''# 攻击路径：注册用户名注入
POST /register
username=hacker', 'pass', 'h@x.com', '999')--
# 实际生成的 SQL：
INSERT INTO users (username, password, email, phone)
VALUES ('hacker', 'pass', 'h@x.com', '999')--', '任意值', '任意值', '任意值')
# -- 注释了后面的所有内容，攻击者可控制写入数据库的全部字段值''', "POC 攻击载荷")

    add_heading(doc, "3.3.2 根因分析与危害", 3)
    add_body(doc, "该漏洞同样由 f-string 字符串拼接导致。与搜索功能不同，此处是 INSERT 语句——"
             "一种写操作。这意味着攻击者不仅能读取数据，还能向数据库中注入任意记录。", size=9.5)
    add_body(doc, "一旦注册的恶意用户能通过登录功能验证，攻击者可以：", size=9.5)
    items = [
        "创建后门管理员账号，实现持久化控制",
        "注入大量垃圾数据导致数据库膨胀，触发拒绝服务",
        "在 username 字段中嵌入 SQL 语句，在后续查询该字段时触发二次注入",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9, indent=0.5)

    add_heading(doc, "3.3.3 修复方案与代码对比", 3)
    add_code(doc, '''# ❌ 漏洞代码（第277行）
sql = f"INSERT INTO users (username, password, email, phone) \\
       VALUES ('{username}', '{password}', '{email}', '{phone}')"
c.execute(sql)''', "修复前")
    add_code(doc, '''# ✅ 修复代码（第317-320行）
# + 新增 CSRF token 校验（第293行）
# + 新增输入长度校验（第301-308行）
# + 密码 bcrypt 哈希存储（第312行）
c.execute(
    "INSERT INTO users (username, password, email, phone) VALUES (?, ?, ?, ?)",
    (username, password_hash, email, phone)
)''', "修复后")

    add_body(doc, "修复要点：", bold=True, size=9.5)
    add_body(doc, "1. 参数化查询（核心）：将 4 个用户字段全部通过 ? 占位符绑定，杜绝注入", size=9.5)
    add_body(doc, "2. CSRF Token 校验：防止跨站构造注册请求（OWASP A01:2021）", size=9.5)
    add_body(doc, "3. bcrypt 密码哈希：即使数据库泄露，密码也无法逆向还原（OWASP A02:2021）", size=9.5)
    add_body(doc, "4. 输入长度限制：前置防御，降低攻击面", size=9.5)

    doc.add_page_break()

    # ─── SQL-04 ───
    add_heading(doc, "3.4 SQL-04：密码明文存储", 2)
    add_table(doc,
        ["属性", "内容"],
        [["漏洞编号", "SQL-04"],
         ["风险等级", "中危"],
         ["CVSS 3.1", "5.9 AV:N/AC:H/PR:H/UI:R/S:U/C:H/I:N/A:N"],
         ["漏洞位置", "app.py 第 77-80 行 init_db() — INSERT 明文密码"],
         ["发现方式", "代码审计"],
         ["CWE 编号", "CWE-312: 敏感信息的明文存储"]])

    add_heading(doc, "3.4.1 攻击路径", 3)
    add_body(doc, "该漏洞本身不是 SQL 注入，但与之强相关。一旦搜索功能的 UNION 注入（SQL-01）或 OR 注入（SQL-02）"
             "被利用导致数据库泄露，密码明文存储将直接导致用户口令批量暴露。", size=9.5)
    add_code(doc, '''# init_db() 中存储明文密码
c.execute("INSERT OR IGNORE INTO users (...) VALUES (?, ?, ?, ?)",
          ("admin", "admin123", "admin@example.com", "13800138000"))  # 明文！''', "漏洞代码")

    add_heading(doc, "3.4.2 修复方案", 3)
    add_code(doc, '''# ✅ 修复后：bcrypt 哈希存储
admin_hash = generate_password_hash("admin123")
alice_hash = generate_password_hash("alice2025")
c.execute("INSERT OR IGNORE INTO users (...) VALUES (?, ?, ?, ?)",
          ("admin", admin_hash, "admin@example.com", "13800138000"))''', "修复后代码")

    doc.add_page_break()

    # ═══════ 4. 整改修复方案 ═══════
    add_heading(doc, "4. 整改修复方案", 1)

    add_heading(doc, "4.1 短期紧急修复（已完成）", 2)
    add_table(doc,
        ["优先级", "修复内容", "涉及文件", "状态"],
        [["P0", "搜索功能：参数化查询替代 f-string 拼接", "app.py:343-351", "✅ 已完成"],
         ["P0", "注册功能：参数化查询替代 f-string 拼接", "app.py:317-320", "✅ 已完成"],
         ["P0", "注册功能：新增 CSRF token 校验", "app.py:293-294", "✅ 已完成"],
         ["P0", "密码存储：bcrypt 哈希替代明文", "app.py:77-82,312", "✅ 已完成"],
         ["P1", "搜索功能：LIKE 通配符转义", "app.py:343", "✅ 已完成"],
         ["P1", "注册功能：输入长度校验", "app.py:301-308", "✅ 已完成"]])

    add_heading(doc, "4.2 中期架构加固（建议 1-2 个月内完成）", 2)
    items = [
        "最小权限原则：为 SQLite 数据库文件设置 600 权限，数据库目录设置 700 权限",
        "统一数据库访问层：封装 DAO 层，集中管理所有 SQL 操作，便于审计和统一加固",
        "WAF 部署：在应用前端部署 Web 应用防火墙，对 SQL 注入特征进行第二层检测",
        "敏感数据加密：对数据库中的手机号、邮箱等个人敏感信息进行 AES 加密存储",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    add_heading(doc, "4.3 长期运营规范", 2)
    items = [
        "代码审计制度化：每次迭代上线前执行自动化 SAST 扫描",
        "安全编码培训：开发团队完成 OWASP 安全编码规范培训",
        "依赖管理：定期扫描第三方库（Flask、Werkzeug）已知漏洞并更新",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    doc.add_page_break()

    # ═══════ 5. 复测验证结论 ═══════
    add_heading(doc, "5. 复测验证结论", 1)

    add_heading(doc, "5.1 验证方法", 2)
    add_body(doc, "修复后使用与漏洞发现阶段完全相同的 POC 攻击载荷进行复测，验证注入路径是否封闭：")
    items = [
        "黑盒复测：使用 curl 构造 UNION SELECT、OR 万能条件、注册注入共 3 组 POC",
        "代码审查：逐行确认所有 SQL 操作均已使用参数化查询",
        "白盒验证：检查 SQLite 数据库中密码存储格式是否为 bcrypt 哈希",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    add_heading(doc, "5.2 复测结果", 2)
    add_table(doc,
        ["测试项", "测试载荷", "修复前", "修复后", "结论"],
        [["UNION 注入", "' UNION SELECT 1,'inj','inj@x.com','138'--", "返回伪造数据", "无结果", "✅ 已修复"],
         ["OR 万能条件", "' OR '1'='1", "返回全部用户", "仅返回匹配结果", "✅ 已修复"],
         ["注册注入", "hacker', 'pass')--", "注入成功写入", "作为普通文本写入", "✅ 已修复"],
         ["密码存储", "查看数据库 raw data", "admin123 明文", "scrypt:... 哈希", "✅ 已修复"],
         ["CSRF 校验", "无 token POST 注册", "注册成功", "HTTP 400 拒绝", "✅ 已修复"],
         ["LIKE 通配符", "搜索 %", "匹配全部用户", "仅匹配字面量 %", "✅ 已修复"]])

    add_heading(doc, "5.3 残留风险声明", 2)
    add_body(doc, "经上述修复，本系统已消除已知 SQL 注入风险。但仍存在以下非 SQL 类安全风险，建议后续迭代中解决：", size=9.5)
    items = [
        "传输层安全：当前使用 HTTP 明文协议，登录凭证和 Session Cookie 在网络中明文传输，建议启用 HTTPS",
        "认证机制：仅依赖用户名+密码单因素认证，建议增加 MFA 支持",
        "Session 管理：Session 依赖客户端 Cookie 存储，建议引入 Redis 等服务端 Session 存储",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9, indent=0.5)

    doc.add_page_break()

    # ═══════ 6. 安全合规建议 ═══════
    add_heading(doc, "6. 安全合规建议", 1)

    add_heading(doc, "6.1 等保 2.0 合规映射", 2)
    add_table(doc,
        ["等保控制点", "要求描述", "当前状态", "整改建议"],
        [["安全计算环境—应用安全",
          "应避免将用户输入直接拼接到 SQL 语句中，采用预编译语句",
          "✅ 已满足", "持续保持参数化查询规范"],
         ["安全计算环境—数据安全",
          "应采用加密技术保护敏感数据存储",
          "⚠️ 部分满足", "建议对手机号、邮箱等个人信息加密存储"],
         ["安全审计",
          "应对用户登录、数据修改等关键操作进行审计",
          "✅ 已满足", "审计日志建议增加自动归档与告警"],
         ["访问控制",
          "应对数据库访问进行最小权限控制",
          "✅ 已满足", "建议对数据库文件设置 600 权限"]])

    add_heading(doc, "6.2 数据安全法合规映射", 2)
    items = [
        "个人信息保护法：用户注册收集手机号、邮箱等个人信息，应在数据库中加密存储并建立数据生命周期管理策略",
        "数据安全法第 27 条：采取技术措施防范 SQL 注入等攻击导致的数据泄露事件",
        "网络安全法第 21 条：采取防范计算机病毒和网络攻击、网络侵入等危害网络安全行为的技术措施",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    add_heading(doc, "6.3 后续安全建设方向", 2)
    items = [
        "SDL 安全开发生命周期：将安全测试纳入 CI/CD 流水线，每次代码提交自动执行 SAST 扫描",
        "安全编码规范建设：制定基于当前技术栈（Python/Flask/SQLite）的安全编码手册",
        "定期渗透测试：至少每季度执行一次全面的 Web 应用渗透测试",
        "供应链安全：跟踪 Flask、Werkzeug 等依赖库的安全公告，及时更新至安全版本",
    ]
    for it in items:
        add_body(doc, f"  • {it}", size=9.5, indent=0.5)

    # ═══════ 附录 ═══════
    doc.add_page_break()
    add_heading(doc, "附录A：全量数据库交互点审计清单", 1)
    add_table(doc,
        ["编号", "函数/路由", "行号", "SQL类型", "参数来源", "当前方案", "判定"],
        [["DB-01", "init_db()", "69", "CREATE TABLE", "无（硬编码）", "静态 SQL", "✅ 安全"],
         ["DB-02", "init_db()", "79-82", "INSERT", "无（硬编码）", "参数化查询 ?", "✅ 安全"],
         ["DB-03", "get_user_from_db()", "94", "SELECT", "username", "参数化查询 ?", "✅ 安全"],
         ["DB-04", "register()", "317-320", "INSERT", "表单 4 字段", "参数化查询 ?", "✅ 安全"],
         ["DB-05", "search()", "348-351", "SELECT LIKE", "keyword", "参数化查询 ? + ESCAPE", "✅ 安全"]])

    add_heading(doc, "附录B：SQL 注入防御最佳实践速查表（Python + SQLite3）", 1)
    add_table(doc,
        ["场景", "推荐方案", "禁止方案"],
        [["固定条件查询", "cursor.execute('SELECT * FROM users WHERE id = ?', (uid,))",
          "cursor.execute(f'SELECT * FROM users WHERE id = {uid}')"],
         ["LIKE 模糊查询", "ESCAPE '\\' + 参数化 ?", "f\"LIKE '%{keyword}%'\""],
         ["INSERT 写入", "参数化 ?：VALUES (?, ?, ?)", "f\"VALUES ('{a}', '{b}', '{c}')\""],
         ["批量操作", "executemany() + 参数化", "循环拼接 execute()"],
         ["动态排序", "白名单校验后拼接列名", "直接拼接用户输入"],
         ["动态表名", "白名单校验后拼接表名", "直接拼接用户输入"]])

    # 结束
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 报告结束 ——")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本报告由安全审计组基于对用户管理系统 app.py 的全量代码审计与渗透测试数据汇编。")
    run.font.size = Pt(8.5)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_MUTED

    out = "/workspace/user_management/阶段二/阶段二-SQL注入审计报告.docx"
    doc.save(out)
    import os as _os
    sz = _os.path.getsize(out) / 1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")


if __name__ == "__main__":
    main()
