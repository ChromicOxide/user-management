#!/usr/bin/env python3
"""SSRF漏洞审查与修复报告"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_PRI = RGBColor(0x1A, 0x2B, 0x3C)
C_SEC = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACC = RGBColor(0x4A, 0x7A, 0xBF)
FONT_HEI = '黑体'; FONT_SONG = '宋体'; FONT_CODE = 'Consolas'
BG_HDR = 'E2E8F0'; BG_ROW = 'F5F7FA'; BG_CODE = 'F0F2F5'

def set_run_font(run, font_name, size=Pt(10), bold=False, color=None):
    run.font.size = size; run.bold = bold
    if color: run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name); rFonts.set(qn('w:hAnsi'), font_name); rFonts.set(qn('w:eastAsia'), font_name)

def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),color)
    c._tc.get_or_add_tcPr().append(s)

def margins(doc, t=2.0, b=2.0, l=2.5, r=2.5):
    for sec in doc.sections: sec.top_margin=Cm(t); sec.bottom_margin=Cm(b); sec.left_margin=Cm(l); sec.right_margin=Cm(r)

def hr(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6)
    bdr=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom'); bt.set(qn('w:val'),'single')
    bt.set(qn('w:sz'),'4'); bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),'D0D5DD'); bdr.append(bt)
    p._element.get_or_add_pPr().append(bdr)

def add_tbl(doc, headers, rows):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    tp=t._tbl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); t._tbl.insert(0,tp)
    cm=OxmlElement('w:tblCellMar')
    for s in ('top','left','bottom','right'):
        m=OxmlElement(f'w:{s}'); m.set(qn('w:w'),'36'); m.set(qn('w:type'),'dxa'); cm.append(m)
    tp.append(cm)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(h); set_run_font(r,FONT_HEI,Pt(8.5),True,C_PRI); shade(c,BG_HDR)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            r=p.add_run(str(val)); set_run_font(r,FONT_SONG,Pt(8),color=C_BODY)
            if ri%2==1: shade(c,BG_ROW)
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(2); sp.paragraph_format.space_after=Pt(2)

def add_code(doc, code, label=""):
    if label:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(1)
        p.paragraph_format.left_indent=Cm(0.3)
        r=p.add_run(label); set_run_font(r,FONT_HEI,Pt(8),True,C_ACC)
    for line in code.rstrip().split('\n'):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.line_spacing=Pt(12)
        r=p.add_run(line); set_run_font(r,FONT_CODE,Pt(7),color=C_BODY)
        sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),BG_CODE)
        p._element.get_or_add_pPr().append(sh)
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(2); sp.paragraph_format.space_after=Pt(2)

def heading(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    sz={1:(18,10,Pt(18)),2:(14,7,Pt(14)),3:(12,5,Pt(12))}.get(level,(10,5,Pt(11)))
    h.paragraph_format.space_before=Pt(sz[0]); h.paragraph_format.space_after=Pt(sz[1])
    clr={1:C_PRI,2:C_SEC,3:C_BODY}.get(level,C_BODY)
    for r in h.runs: set_run_font(r,FONT_HEI,sz[2],True,clr)
    if level==1: hr(doc)

def body(doc, text, bold=False, size=10, indent=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); set_run_font(r,FONT_SONG,Pt(size),bold,color=C_BODY)

def set_style(doc):
    s=doc.styles['Normal']; s.font.name=FONT_SONG; s.font.size=Pt(10)
    s.paragraph_format.line_spacing=1.35; s.paragraph_format.space_after=Pt(4); s.font.color.rgb=C_BODY
    rPr=s.element.get_or_add_rPr()
    rFonts=rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts=OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'),FONT_SONG); rFonts.set(qn('w:ascii'),FONT_SONG); rFonts.set(qn('w:hAnsi'),FONT_SONG)

def main():
    doc=Document(); set_style(doc); margins(doc)

    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("SSRF 服务器端请求伪造\n漏洞审查与修复报告"); r.bold=True; r.font.size=Pt(24); r.font.name=FONT_HEI; r.font.color.rgb=C_PRI
    doc.add_paragraph()
    info=[("文档编号","SEC-SSRF-2025-001"),("密级","内部公开"),("版本号","V1.0"),("发布日期","2026-07-15"),("编制人","潘麒宇")]
    tbl=doc.add_table(rows=len(info), cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(info):
        for j,t in enumerate([k,v]):
            c=tbl.rows[i].cells[j]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(t); r.font.size=Pt(11); r.font.name=FONT_SONG
            if j==0: r.bold=True; r.font.color.rgb=C_PRI
            else: r.font.color.rgb=C_ACC
            shade(c,'EDF2F7' if i%2==0 else 'F5F8FB')
    doc.add_page_break()

    heading(doc,"1. 审计概述",1)
    heading(doc,"1.1 项目信息",2)
    add_tbl(doc,["项目","内容"],[
        ["系统名称","用户管理系统"],
        ["审计范围","URL 抓取功能（/fetch-url 路由 + urllib 请求处理）"],
        ["审计日期","2026 年 7 月 15 日"],
        ["技术栈","Python 3 / Flask 3.x / urllib"],
    ])
    heading(doc,"1.2 审计方法",2)
    for it in ["代码审计：逐行审查 fetch_url 路由的逻辑，定位 SSRF 利用路径",
               "黑盒渗透：构造 file:// 协议、内网 IP、云元数据等攻击载荷验证 SSRF 存在",
               "参考标准：OWASP Top 10 A10:2021 / CWE-918 / CWE-200"]:
        body(doc,f"  * {it}",size=9.5,indent=0.5)

    doc.add_page_break()
    heading(doc,"2. 漏洞发现清单",1)
    heading(doc,"2.1 漏洞总览",2)
    add_tbl(doc,["编号","漏洞名称","等级","所在路由","漏洞类型"],[
        ["SSRF-01","无协议白名单 — file:// 读取任意文件","高危","POST /fetch-url","SSRF/文件读取"],
        ["SSRF-02","无内网 IP 限制 — 内网服务探测","高危","POST /fetch-url","SSRF/内网扫描"],
        ["SSRF-03","云元数据攻击 — 云主机凭证泄露","高危","POST /fetch-url","SSRF/云攻击"],
        ["SSRF-04","无端口限制 — 内网端口扫描","中危","POST /fetch-url","SSRF/信息收集"],
        ["SSRF-05","无返回值限制 — 响应内容泄露","低危","POST /fetch-url","信息泄露"],
        ["SSRF-06","无 URL 长度限制 — 日志注入","低危","POST /fetch-url","日志攻击"],
        ["SSRF-07","无重定向限制 — 302跳转绕过","中危","POST /fetch-url","SSRF绕过"],
    ])
    heading(doc,"2.2 风险分布",2)
    add_tbl(doc,["风险等级","数量","占比"],[["高危","3个","42.9%"],["中危","2个","28.6%"],["低危","2个","28.6%"],["合计","7个","100%"]])

    doc.add_page_break()
    heading(doc,"3. 单漏洞深度分析与修复",1)

    vulns=[
        ("SSRF-01","无协议白名单 — file:// 读取任意文件","高危",
         "app.py:672 fetch_url() — urllib.request.Request(target_url)",
         "攻击路径：\n1. 登录后访问首页 URL 抓取功能\n2. 输入 file:///etc/passwd\n3. 点击抓取\n4. 系统密码文件内容返回到页面上\n\n攻击载荷：\nfile:///etc/passwd     → 读取系统密码\nfile:///app/app.py     → 读取源码\nfile:///proc/1/environ → 读取环境变量",
         "CVSS 3.1: 9.1 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "攻击者可读取服务器上任意文件，包括源码、配置文件、系统密码文件等，导致全量信息泄露。",
         "修复方案：\n1. 协议白名单：仅允许 http:// 和 https://\n2. 校验 parsed.scheme 是否在 ALLOWED_PROTOCOLS 中\n3. 不在白名单中返回 不支持的协议类型"),
        ("SSRF-02","无内网 IP 限制 — 内网服务探测","高危",
         "app.py:672 fetch_url() — 无目标 IP 校验",
         "攻击路径：\n1. 提交 http://127.0.0.1:5000/ 抓取自身\n2. 提交 http://127.0.0.1:3306 探测 MySQL\n3. 提交 http://127.0.0.1:6379 探测 Redis\n4. 遍历内网网段 10.x.x.x:8080 等",
         "CVSS 3.1: 8.6 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "攻击者可利用服务器作为跳板，探测内网未授权服务。",
         "修复方案：\n1. 引入 _is_internal_ip() 函数\n2. 禁止 127.0.0.0/8、10.0.0.0/8、172.16.0.0/12、192.168.0.0/16\n3. 禁止 localhost 等域名"),
        ("SSRF-03","云元数据攻击","高危",
         "app.py:672 fetch_url() — 未过滤云元数据 IP",
         "攻击路径：\n1. http://169.254.169.254/latest/meta-data/\n2. http://100.100.100.200（阿里云）\n3. 获取云主机临时凭证、SSH 公钥等",
         "CVSS 3.1: 9.1 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H",
         "云环境下可获取 IAM 临时凭证，导致整个云账号沦陷。",
         "修复方案：将 169.254.169.254 和 100.100.100.200 加入 FORBIDDEN_IPS"),
        ("SSRF-04","无端口限制","中危",
         "app.py:672 — 不检查目标端口",
         "攻击路径：\n1. http://127.0.0.1:22 → 探测 SSH\n2. http://127.0.0.1:3306 → 探测 MySQL\n3. http://127.0.0.1:6379 → 探测 Redis",
         "CVSS 3.1: 5.3", "内网端口扫描，辅助横向移动。",
         "修复方案：已通过 SSRF-02 内网 IP 拦截一并解决。"),
        ("SSRF-05","无返回值限制","低危",
         "app.py:676 — 响应内容返回给用户",
         "响应内容直接呈现在页面上，内网服务的敏感数据被泄露。",
         "CVSS 3.1: 3.1", "内网服务返回的数据被展示在页面上。",
         "修复方案：已有 content[:5000] 长度限制，无需额外操作。"),
        ("SSRF-06","无 URL 长度限制","低危",
         "app.py:665 — 无长度校验",
         "超长 URL 可能导致日志文件被注入恶意内容，或服务拒绝。",
         "CVSS 3.1: 3.3", "日志注入、资源消耗。",
         "修复方案：添加 len(target_url) > 2048 限制。"),
        ("SSRF-07","无重定向限制","中危",
         "app.py:673 — urllib 默认跟随 302",
         "攻击路径：\n1. 外部服务器设置 302 跳转到 file:///etc/passwd\n2. urllib 跟随跳转读取文件",
         "CVSS 3.1: 5.8", "通过 302 跳转绕过协议白名单限制。",
         "修复方案：自定义 NoRedirectHandler 替代默认 opener"),
    ]

    for idx,(vid,vname,vlevel,vloc,vpath,vcvss,vharm,vfix) in enumerate(vulns):
        heading(doc,f"3.{idx+1} {vid}：{vname} [{vlevel}]",2)
        add_tbl(doc,["属性","内容"],[
            ["漏洞编号",vid],["风险等级",vlevel],["代码位置",vloc],
            ["CVSS 3.1",vcvss],["危害程度",vharm]])
        add_code(doc,vpath,"攻击路径与根因分析")
        add_code(doc,vfix,"修复方案")

    doc.add_page_break()
    heading(doc,"4. 修复后核心代码",1)
    add_code(doc,
        "ALLOWED_PROTOCOLS = {\"http\", \"https\"}\n\n"
        "def _is_internal_ip(host: str) -> bool:\n"
        "    if host.lower() in {\"127.0.0.1\", \"localhost\", \"169.254.169.254\", ...}:\n"
        "        return True\n"
        "    for ip in socket.getaddrinfo(host, None):\n"
        "        obj = ipaddress.ip_address(ip[4][0])\n"
        "        for net in [\"127.0.0.0/8\", \"10.0.0.0/8\", \"172.16.0.0/12\", \"192.168.0.0/16\"]:\n"
        "            if obj in ipaddress.ip_network(net): return True\n"
        "    return False\n\n"
        "parsed = urlparse(target_url)\n"
        "if parsed.scheme not in ALLOWED_PROTOCOLS:          # 协议白名单\n"
        "    return \"不支持的协议类型\"\n"
        "host = parsed.hostname\n"
        "if _is_internal_ip(host):                           # 内网检测\n"
        "    return \"不允许访问内网地址\"\n\n"
        "# 自定义 opener 禁止跳转\n"
        "class NoRedirectHandler(HTTPRedirectHandler):\n"
        "    def redirect_request(self, req, fp, code, msg, headers, newurl):\n"
        "        return None",
        "修复后 /fetch-url 路由核心安全逻辑")

    doc.add_page_break()
    heading(doc,"5. 修复验证结论",1)
    add_tbl(doc,["测试项","测试 URL","修复前","修复后","结论"],[
        ["正常 HTTP","http://example.com","返回200","返回200","✓"],
        ["HTTPS","https://baidu.com","返回200","返回200","✓"],
        ["file 协议","file:///etc/passwd","读取文件","不支持协议","✓"],
        ["内网 127.0.0.1","http://127.0.0.1:5000/","访问成功","内网拦截","✓"],
        ["内网 localhost","http://localhost:5000/","访问成功","内网拦截","✓"],
        ["云元数据","http://169.254.169.254/","可访问","内网拦截","✓"],
        ["内网 10.x","http://10.0.0.1/","可访问","内网拦截","✓"],
        ["内网 192.168","http://192.168.1.1/","可访问","内网拦截","✓"],
        ["内网 172.16","http://172.16.0.1/","可访问","内网拦截","✓"],
    ])

    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("—— 报告结束 ——"); r.bold=True; r.font.size=Pt(14); r.font.name=FONT_HEI; r.font.color.rgb=C_PRI
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("本报告由安全审计组基于对用户管理系统 URL 抓取功能的代码审计与渗透测试数据汇编。")
    r.font.size=Pt(8.5); r.font.name=FONT_SONG; r.font.color.rgb=C_MUTED

    out="/workspace/user_management/SEC-SSRF-AUDIT-2025-001.docx"
    doc.save(out)
    import os; sz=os.path.getsize(out)/1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")

if __name__=="__main__":
    main()
