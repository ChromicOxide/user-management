#!/usr/bin/env python3
"""权限提升与业务逻辑漏洞审查与修复报告 — 字体修复版"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 配色 ──
C_PRI = RGBColor(0x1A, 0x2B, 0x3C)
C_SEC = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACC = RGBColor(0x4A, 0x7A, 0xBF)

# ── 字体 ──
FONT_HEI = '黑体'
FONT_SONG = '宋体'
FONT_CODE = 'Consolas'

BG_HDR = 'E2E8F0'
BG_ROW = 'F5F7FA'
BG_CODE = 'F0F2F5'


def set_run_font(run, font_name, size=Pt(10), bold=False, color=None):
    """设置 run 的中英文字体：分别指定 ASCII、EastAsia、HAnsi"""
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def shade(c, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), color)
    c._tc.get_or_add_tcPr().append(s)


def margins(doc, t=2.0, b=2.0, l=2.5, r=2.5):
    for sec in doc.sections:
        sec.top_margin = Cm(t)
        sec.bottom_margin = Cm(b)
        sec.left_margin = Cm(l)
        sec.right_margin = Cm(r)


def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    bdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single')
    bt.set(qn('w:sz'), '4')
    bt.set(qn('w:space'), '1')
    bt.set(qn('w:color'), 'D0D5DD')
    bdr.append(bt)
    p._element.get_or_add_pPr().append(bdr)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    tp = t._tbl.tblPr
    if tp is None:
        tp = OxmlElement('w:tblPr')
        t._tbl.insert(0, tp)
    cm = OxmlElement('w:tblCellMar')
    for s in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement(f'w:{s}')
        m.set(qn('w:w'), '36')
        m.set(qn('w:type'), 'dxa')
        cm.append(m)
    tp.append(cm)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h)
        set_run_font(r, FONT_HEI, Pt(8.5), bold=True, color=C_PRI)
        shade(c, BG_HDR)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val))
            set_run_font(r, FONT_SONG, Pt(8), color=C_BODY)
            if ri % 2 == 1:
                shade(c, BG_ROW)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)


def add_code(doc, code, label=""):
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(label)
        set_run_font(r, FONT_HEI, Pt(8), bold=True, color=C_ACC)
    for line in code.rstrip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = Pt(12)
        r = p.add_run(line)
        set_run_font(r, FONT_CODE, Pt(7), color=C_BODY)
        sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear')
        sh.set(qn('w:color'), 'auto')
        sh.set(qn('w:fill'), BG_CODE)
        p._element.get_or_add_pPr().append(sh)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)


def heading_styled(doc, text, level=1):
    """使用 add_heading 后覆盖字体设置，确保中英文均使用黑体"""
    h = doc.add_heading(text, level=level)
    sz_map = {1: (18, 10, Pt(18)), 2: (14, 7, Pt(14)), 3: (12, 5, Pt(12))}
    sb, sa, sz = sz_map.get(level, (10, 5, Pt(11)))
    h.paragraph_format.space_before = Pt(sb)
    h.paragraph_format.space_after = Pt(sa)
    color_map = {1: C_PRI, 2: C_SEC, 3: C_BODY}
    clr = color_map.get(level, C_BODY)
    for r in h.runs:
        set_run_font(r, FONT_HEI, sz, bold=True, color=clr)
    # 同时修改底层的样式定义，保证 Word 渲染不覆盖
    style_name = f'Heading {level}'
    if style_name in doc.styles:
        sty = doc.styles[style_name]
        sty.font.name = FONT_HEI
        rPr = sty.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), FONT_HEI)
        rFonts.set(qn('w:hAnsi'), FONT_HEI)
        rFonts.set(qn('w:eastAsia'), FONT_HEI)
    if level == 1:
        hr(doc)
    return h


def body(doc, text, bold=False, size=10, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_run_font(r, FONT_SONG, Pt(size), bold=bold, color=C_BODY)


def set_default_style(doc):
    """修改 Normal 样式为宋体"""
    s = doc.styles['Normal']
    s.font.name = FONT_SONG
    s.font.size = Pt(10)
    s.paragraph_format.line_spacing = 1.35
    s.paragraph_format.space_after = Pt(4)
    s.font.color.rgb = C_BODY
    # 设置 east-asian 字体
    rPr = s.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_SONG)
    rFonts.set(qn('w:ascii'), FONT_SONG)
    rFonts.set(qn('w:hAnsi'), FONT_SONG)


def main():
    doc = Document()
    set_default_style(doc)
    margins(doc)

    # ═══ 封面 ═══
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("权限提升与业务逻辑漏洞\n审查与修复报告")
    set_run_font(r, FONT_HEI, Pt(24), bold=True, color=C_PRI)
    doc.add_paragraph()

    info = [("文档编号", "SEC-PRIV-2025-001"),
            ("密级", "内部公开"),
            ("版本号", "V1.0"),
            ("发布日期", "2026-07-09"),
            ("编制人", "潘麒宇")]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        for j, t in enumerate([k, v]):
            c = tbl.rows[i].cells[j]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(t)
            if j == 0:
                set_run_font(r, FONT_SONG, Pt(11), bold=True, color=C_PRI)
            else:
                set_run_font(r, FONT_SONG, Pt(11), color=C_ACC)
            shade(c, 'EDF2F7' if i % 2 == 0 else 'F5F8FB')
    doc.add_page_break()

    # ═══ 1 ═══
    heading_styled(doc, "1. 审计概述", 1)
    heading_styled(doc, "1.1 项目信息", 2)
    add_table(doc, ["项目", "内容"], [
        ["系统名称", "用户管理系统"],
        ["审计范围", "权限校验逻辑、充值业务逻辑、用户身份验证全链路"],
        ["技术栈", "Python 3 / Flask 3.x / SQLite3"],
    ])
    heading_styled(doc, "1.2 审计方法", 2)
    for it in ["黑盒渗透：Burp Suite 拦截并篡改请求参数，验证越权访问路径",
               "代码审计：逐行审查身份校验逻辑、权限判定条件",
               "业务逻辑推演：充值的正负数处理、余额上限、IDOR 链路"]:
        body(doc, f"  * {it}", size=9.5, indent=0.5)
    heading_styled(doc, "1.3 参考标准", 2)
    add_table(doc, ["标准", "适用范围"], [
        ["OWASP Top 10 2021", "A01:2021 访问控制失效"],
        ["CWE-639", "通过用户控制的键绕过授权"],
        ["CWE-840", "业务逻辑错误"],
    ])
    doc.add_page_break()

    # ═══ 2 ═══
    heading_styled(doc, "2. 漏洞发现清单", 1)
    heading_styled(doc, "2.1 漏洞总览", 2)
    add_table(doc, ["编号", "漏洞名称", "等级", "所在路由", "漏洞类型"], [
        ["P-01", "IDOR 越权查看任意用户", "高危", "GET /profile?user_id=X", "权限提升"],
        ["P-02", "个人中心无需登录", "高危", "GET /profile", "权限提升"],
        ["P-03", "负数充值任意扣减余额", "高危", "POST /recharge", "业务逻辑"],
        ["P-04", "充值无归属校验", "中危", "POST /recharge", "业务逻辑"],
        ["P-05", "余额无上限溢出", "低危", "POST /recharge", "业务逻辑"],
        ["P-06", "用户 ID 枚举遍历", "低危", "GET /profile?user_id=X", "信息泄露"],
        ["P-07", "权限分级形同虚设", "中危", "全局", "权限提升"],
    ])
    heading_styled(doc, "2.2 风险分布", 2)
    add_table(doc, ["风险等级", "数量", "占比"], [
        ["高危", "3个", "42.9%"],
        ["中危", "2个", "28.6%"],
        ["低危", "2个", "28.6%"],
        ["合计", "7个", "100%"],
    ])
    doc.add_page_break()

    # ═══ 3 ═══
    heading_styled(doc, "3. 单漏洞深度分析与修复", 1)

    vulns_data = [
        ("P-01", "IDOR 越权查看任意用户", "高危",
         "app.py:518-529 profile() 路由",
         "攻击路径：\n1. 登录 alice → 导航栏写死 user_id=1\n2. 点击个人中心 → 看到 admin 资料\n3. 修改 URL 为 ?user_id=3,4,5... 遍历所有用户\n\n代码根因：\n@app.route(\"/profile\")       # 无 @login_required\ndef profile():\n    user_id = request.args.get(\"user_id\", type=int)\n    user_data = get_user_by_id(user_id)\n    return render_template(...)",
         "CVSS 3.1: 8.6 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "攻击者可遍历 user_id 获取系统全部注册用户的邮箱、手机号、余额等敏感信息。",
         "修复：1. @login_required 2. session 与 user_id 比对 3. nav 链接改为 /profile"),
        ("P-02", "个人中心无需登录", "高危",
         "app.py:518 profile() 无 @login_required",
         "攻击路径：\n1. 清除 Cookie\n2. 直接访问 /profile?user_id=1\n3. 返回 admin 完整资料",
         "CVSS 3.1: 7.5 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "完全未认证即可获取全量用户数据。",
         "修复：添加 @login_required 装饰器"),
        ("P-03", "负数充值任意扣减余额", "高危",
         "app.py:532-537 recharge() 未校验正负",
         "攻击路径：\n1. 登录任意账号\n2. Burp 拦截充值 POST\n3. amount=100 -> amount=-99999\n4. Forward -> 余额减少",
         "CVSS 3.1: 7.5 | AV:N/AC:L/PR:L/UI:N/S:U/C:N/I:H/A:N",
         "攻击者可耗尽任意用户的余额。",
         "修复：校验 amount > 0，单次上限 1000000，MAX_BALANCE"),
        ("P-04", "充值无归属校验", "中危",
         "app.py:532-537 recharge() 未校验当前用户",
         "攻击路径：\n1. 登录 alice\n2. Burp 拦截充值请求\n3. user_id=2 -> user_id=1\n4. Forward -> alice 给 admin 充了值",
         "CVSS 3.1: 6.5 | AV:N/AC:L/PR:L/UI:N/S:U/C:N/I:L/A:N",
         "任何人可操控他人账户余额。",
         "修复：充值时 user_id 与 session 当前用户 ID 比对"),
        ("P-05", "余额无上限溢出", "低危",
         "app.py:532-537 recharge() update_balance() 无上限校验",
         "攻击路径：\n1. 登录任意账号\n2. Burp 拦截充值 POST 请求\n3. amount=100 改为 amount=999999999999999\n4. Forward 放行\n5. 余额变为异常值（整数溢出/无限大）",
         "CVSS 3.1: 3.3 | AV:N/AC:L/PR:L/UI:N/S:U/C:N/I:N/A:L",
         "整数溢出可能导致业务逻辑异常，余额字段失控。",
         "修复：设置 MAX_BALANCE=999999999，单次充值上限 1000000"),
        ("P-06", "用户 ID 枚举遍历", "低危",
         "app.py:518-529 profile() 未限制 user_id 范围",
         "攻击路径：\n1. 通过 Burp Repeater 遍历 user_id 参数\n2. user_id=1 → admin\n3. user_id=2 → alice\n4. user_id=3,4,5... → 遍历所有注册用户\n5. 可获取系统全部用户数量及信息",
         "CVSS 3.1: 3.3 | AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N",
         "暴露系统用户数量及全量用户信息，辅助社工攻击。",
         "修复：已通过 P-01 IDOR 修复一并解决，非当前用户统一返回无权查看"),
        ("P-07", "权限分级形同虚设", "中危",
         "全局，USERS 字典中 admin 有 role=admin 但全系统未校验",
         "攻击路径：\n1. 注册任意普通用户（/register）\n2. 发现所有功能与 admin 完全一致\n3. 搜索功能可查全量用户\n4. 上传/充值功能均可正常使用\n5. role 字段仅为标记，无实际权限控制效力",
         "CVSS 3.1: 5.3 | AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N",
         "不存在实际权限分级，admin 与普通用户权限无区别，违反最小权限原则。",
         "建议：后续引入 RBAC 模型，基于 role 做路由级权限校验。当前仅记录，未改动已有代码。"),
    ]

    for idx, (vid, vname, vlevel, vloc, vpath, vcvss, vharm, vfix) in enumerate(vulns_data):
        heading_styled(doc, f"3.{idx+1} {vid}：{vname} [{vlevel}]", 2)
        add_table(doc, ["属性", "内容"], [
            ["漏洞编号", vid],
            ["风险等级", vlevel],
            ["代码位置", vloc],
            ["CVSS 3.1", vcvss],
            ["危害程度", vharm],
        ])
        add_code(doc, vpath, "攻击路径与根因分析")
        body(doc, vfix, size=9.5)

    doc.add_page_break()

    # ═══ 4 ═══
    heading_styled(doc, "4. 修复后核心代码", 1)
    add_code(doc,
        '@app.route("/profile")\n'
        '@login_required\n'
        'def profile():\n'
        '    current_user = session.get("username")\n'
        '    current_id = get_user_id_by_username(current_user)\n'
        '    req_id = request.args.get("user_id", type=int)\n'
        '    target_id = req_id if req_id else current_id\n'
        '    if target_id != current_id:\n'
        '        return render_template("profile.html", user_data=None,\n'
        '                               error="无权查看其他用户的资料")\n'
        '    user_data = get_user_by_id(target_id)\n'
        '    return render_template("profile.html", user_data=user_data)',
        "修复后 /profile 路由")
    add_code(doc,
        '@app.route("/recharge", methods=["POST"])\n'
        '@login_required\n'
        'def recharge():\n'
        '    current_id = get_user_id_by_username(session.get("username"))\n'
        '    req_id = request.form.get("user_id", type=int)\n'
        '    amount = request.form.get("amount", type=int, default=0)\n'
        '    if req_id != current_id:\n'
        '        return redirect("/")\n'
        '    if amount <= 0:\n'
        '        return redirect("/profile?user_id=" + str(req_id))\n'
        '    if amount > 1000000:\n'
        '        amount = 1000000\n'
        '    update_balance(req_id, amount)\n'
        '    return redirect("/profile?user_id=" + str(req_id))',
        "修复后 /recharge 路由")

    doc.add_page_break()

    # ═══ 5 ═══
    heading_styled(doc, "5. 修复验证结论", 1)
    add_table(doc, ["测试项", "操作方式", "修复前", "修复后", "结论"], [
        ["未登录访问 profile", "清除 Cookie -> GET /profile", "返回用户数据", "401 拒绝", "✓"],
        ["IDOR 越权查 admin", "alice 登录 -> user_id=1", "返回 admin 资料", "无权查看", "✓"],
        ["个人中心自己", "alice 登录 -> /profile", "N/A", "正常显示", "✓"],
        ["负数充值", "amount=-500", "余额减少", "不处理负数", "✓"],
        ["越权充值", "user_id=1 (alice 操作)", "admin 余额变动", "跳转首页", "✓"],
        ["整数溢出", "amount=9999999999", "可能异常", "上限拦截", "✓"],
        ["枚举遍历", "user_id=1,2,3...", "遍历出用户列表", "仅返回本人", "✓"],
    ])
    heading_styled(doc, "6. 后续建议", 1)
    for it in ["引入 RBAC 权限模型，为不同角色分配不同操作权限",
               "日志告警：频繁 IDOR 扫描行为触发告警",
               "API 统一鉴权：敏感接口经过鉴权中间件"]:
        body(doc, f"  * {it}", size=9.5, indent=0.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 报告结束 ——")
    set_run_font(r, FONT_HEI, Pt(14), bold=True, color=C_PRI)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本报告由安全审计组基于对用户管理系统的全量权限审计数据汇编。")
    set_run_font(r, FONT_SONG, Pt(8.5), color=C_MUTED)

    out = "/workspace/user_management/阶段四/阶段四-权限提升审计报告.docx"
    doc.save(out)
    import os
    sz = os.path.getsize(out) / 1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")


if __name__ == "__main__":
    main()

# 39152523
