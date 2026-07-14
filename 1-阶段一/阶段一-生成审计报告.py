#!/usr/bin/env python3
"""生成安全审计报告 Word 文档 — 优化字体与排版"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import datetime


# ── 调色板 ──────────────────────────────────────────────
C_PRIMARY   = RGBColor(0x1A, 0x2B, 0x3C)   # 深蓝灰 — 主标题/表头字
C_SECONDARY = RGBColor(0x4A, 0x5A, 0x6A)   # 中灰 — 二级标题
C_BODY      = RGBColor(0x2C, 0x2C, 0x2C)   # 正文
C_MUTED     = RGBColor(0x8C, 0x8C, 0x8C)   # 辅助信息
C_ACCENT    = RGBColor(0x4A, 0x7A, 0xBF)   # 强调色
BG_HEADER   = 'E2E8F0'                      # 表头底色
BG_ROW      = 'F5F7FA'                      # 行交替色
BG_CODE     = 'F0F2F5'                      # 代码块底色


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """设置 A4 页面边距"""
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_horizontal_rule(doc):
    """在段落下方插入水平线（模拟 Word 下边框）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D0D5DD')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加表格 — 浅色紧凑"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 缩小单元格内边距
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '36')
        m.set(qn('w:type'), 'dxa')
        cell_mar.append(m)
    tblPr.append(cell_mar)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.name = '微软雅黑'
        run.font.color.rgb = C_PRIMARY
        set_cell_shading(cell, BG_HEADER)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = '微软雅黑'
            run.font.color.rgb = C_BODY
            if ri % 2 == 1:
                set_cell_shading(cell, BG_ROW)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # 表后间距
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after  = Pt(2)
    spacer.paragraph_format.line_spacing = Pt(8)
    return table


def add_code_block(doc, code_text, label=""):
    """添加代码块 — 等宽字体 + 灰底"""
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(f"▎{label}")
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = '微软雅黑'
        run.font.color.rgb = C_ACCENT

    for line in code_text.rstrip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name  = 'Consolas'
        run.font.size  = Pt(7)
        run.font.color.rgb = C_BODY
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), BG_CODE)
        p._element.get_or_add_pPr().append(shd)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after  = Pt(2)
    spacer.paragraph_format.line_spacing = Pt(8)


def add_heading_styled(doc, text, level=1):
    """标题 — 层级清晰"""
    h = doc.add_heading(text, level=level)
    spacing = {1: (16, 8), 2: (10, 4)}.get(level, (8, 3))
    h.paragraph_format.space_before = Pt(spacing[0])
    h.paragraph_format.space_after  = Pt(spacing[1])
    for run in h.runs:
        run.font.name = '微软雅黑'
        if level == 1:
            run.font.size  = Pt(16)
            run.font.color.rgb = C_PRIMARY
        elif level == 2:
            run.font.size  = Pt(13)
            run.font.color.rgb = C_SECONDARY
        else:
            run.font.size  = Pt(11)
            run.font.color.rgb = C_BODY
    # 一级标题下方加分隔线
    if level == 1:
        add_horizontal_rule(doc)
    return h


def add_body(doc, text, bold=False, size=10):
    """正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.color.rgb = C_BODY
    run.bold = bold
    return p


def add_bullet(doc, text, indent=0.5):
    """要点列表"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_BODY
    return p


def set_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after  = Pt(4)
    style.font.color.rgb = C_BODY


# ═══════════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════════
def main():
    doc = Document()
    set_default_style(doc)
    set_page_margins(doc)

    # ── 封面 ──────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("用户管理系统\n未登录页面安全审计报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_PRIMARY

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)

    info = [
        ("文档编号", "SEC-AUDIT-2025-001"),
        ("密  级", "内部公开"),
        ("版本号", "V2.0"),
        ("发布日期", "2026-07-07"),
        ("编制人", "潘麒宇"),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        for j, txt in enumerate([k, v]):
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt)
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            if j == 0:
                run.bold = True
                run.font.color.rgb = C_PRIMARY
            else:
                run.font.color.rgb = C_ACCENT
            set_cell_shading(cell, 'EDF2F7' if i % 2 == 0 else 'F5F8FB')
    doc.add_page_break()

    # ── 目录 ──────────────────────────────────────────
    add_heading_styled(doc, "目录", 1)
    toc = [
        ("1", "审计基本信息概览"),
        ("2", "初始代码漏洞扫描结果"),
        ("3", "漏洞整改过程说明"),
        ("4", "整改前后代码对比"),
        ("5", "最终整改结果验证"),
        ("6", "总结与后续建议"),
        ("附录A", "文件清单"),
        ("附录B", "修复数据汇总"),
    ]
    for num, title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title}")
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
        run.font.color.rgb = C_PRIMARY if num.isdigit() else C_SECONDARY

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 1. 审计基本信息概览
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "1. 审计基本信息概览", 1)

    add_heading_styled(doc, "1.1 审计对象", 2)
    add_table_with_style(doc,
        ["项目", "内容"],
        [["系统名称", "用户管理系统（User Management System）"],
         ["审计范围", "未登录页面（/login、/、base.html）及后端鉴权逻辑"],
         ["代码仓库", "/workspace/user_management/"],
         ["技术栈", "Python Flask 3.x + Werkzeug 3.x + Jinja2 + HTML5/CSS3"]])

    add_heading_styled(doc, "1.2 审计时间", 2)
    add_table_with_style(doc,
        ["阶段", "时间", "内容"],
        [["第一轮审计", "2026-07-07", "初始代码全量漏洞扫描"],
         ["第二轮修复", "2026-07-07", "高危/中危漏洞集中修复"],
         ["第三轮加固", "2026-07-07", "残余漏洞修复 + 安全纵深"],
         ["合规审计", "2026-07-07", "OWASP Top10 + WCAG 2.1 合规验收"]])

    add_heading_styled(doc, "1.3 测试环境", 2)
    add_table_with_style(doc,
        ["环境", "配置"],
        [["服务端", "Kali Linux, Python 3.13.12, Flask 3.x"],
         ["客户端", "Windows 10/11, Chrome/Firefox 最新版"],
         ["IP地址", "192.168.72.129:5000（内网测试）"],
         ["通信协议", "HTTP（开发环境）"]])

    add_heading_styled(doc, "1.4 参考标准", 2)
    add_table_with_style(doc,
        ["标准", "版本", "适用范围"],
        [["OWASP Top 10", "2024", "Web应用安全风险分类与评级"],
         ["WCAG", "2.1 AA", "Web内容无障碍访问标准"],
         ["ISO 27001", "2022", "信息安全控制措施"],
         ["《信息安全等级保护》", "基本要求(2024)", "安全通用技术要求"]])

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 2. 初始代码漏洞扫描结果
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "2. 初始代码漏洞扫描结果", 1)

    add_heading_styled(doc, "2.1 初始代码概况", 2)
    add_body(doc, "审计人员接收到的初始代码为一份未登录页面的HTML源码，其核心实现包含：")
    for item in [
        "后端（app.py）：Flask 应用，密码明文存储于字典，secret_key 硬编码为 dev-key-2025，debug=True",
        "前端（login.html）：HTML注释中明文写入管理员账号密码",
        "前端（index.html）：登录成功后直接展示包含密码字段在内的全部用户信息",
        "样式（style.css）：基础UI样式，无安全考量",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "2.2 漏洞清单总表", 2)
    add_body(doc, "经三轮逐层审计，共计发现 19 个安全漏洞，按 OWASP 分类列示如下：")
    vulns = [
        ["V-01", "A5", "密码明文存储在数据库字典中", "高危", "app.py USERS字典", "数据库泄露即密码全量暴露"],
        ["V-02", "A5", "secret_key 硬编码为弱密钥", "高危", "app.py L4", "攻击者可伪造任意session"],
        ["V-03", "A5", "Flask Debug模式开启", "高危", "app.py debug=True", "远程代码执行（RCE）"],
        ["V-04", "A1", "HTML注释泄露管理员账号密码", "高危", "login.html L1", "查看源代码即可获取admin/admin123"],
        ["V-05", "A4", "登录后密码明文展示在首页", "高危", "index.html L8", "敏感信息直接暴露给UI用户"],
        ["V-06", "A7", "密码明文比对（==比较）", "中危", "app.py login路由", "无哈希保护，时序可测"],
        ["V-07", "A7", "CSRF防护缺失", "中危", "login.html 表单", "跨站请求伪造攻击"],
        ["V-08", "A7", "暴力破解无防护", "中危", "app.py login路由", "可无限次猜测密码"],
        ["V-09", "A7", "会话固定攻击", "中危", "app.py session管理", "登录前后session ID不变"],
        ["V-10", "A5", "Session Cookie无安全标记", "中危", "app.py cookie配置", "HttpOnly/SameSite未设置"],
        ["V-11", "A5", "Server头泄露版本号", "中危", "Werkzeug默认行为", "暴露Werkzeug/Python精确版本"],
        ["V-12", "A7", "登录态无过期时间", "中危", "app.py session配置", "session永不过期"],
        ["V-13", "A7", "用户名枚举（时序侧信道）", "中危", "app.py login路由", "测响应时间可判断用户是否存在"],
        ["V-14", "A3", "输入校验缺失", "低危", "login.html 表单", "无长度限制"],
        ["V-15", "A5", "CSP缺失", "低危", "base.html", "页面可被XSS注入"],
        ["V-16", "A5", "X-Content-Type-Options缺失", "低危", "HTTP响应头", "MIME类型嗅探攻击"],
        ["V-17", "A5", "X-Frame-Options缺失", "低危", "HTTP响应头", "点击劫持"],
        ["V-18", "A9", "审计日志缺失", "低危", "app.py 全局", "安全事件不可追溯"],
        ["V-19", "A5", "LOGIN_ATTEMPTS无大小限制", "低危", "app.py 全局", "内存无限膨胀导致DoS"],
    ]
    add_table_with_style(doc,
        ["编号", "OWASP", "漏洞名称", "等级", "位置", "风险影响"], vulns)

    add_heading_styled(doc, "2.3 漏洞风险统计", 2)
    add_table_with_style(doc,
        ["风险等级", "数量", "占比"],
        [["高危", "5个", "26.3%"],
         ["中危", "8个", "42.1%"],
         ["低危", "6个", "31.6%"],
         ["合计", "19个", "100%"]])

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 3. 漏洞整改过程说明
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "3. 漏洞整改过程说明", 1)

    add_heading_styled(doc, "3.1 整体修复策略", 2)
    add_body(doc, "本次整改采用三轮递进式修复策略，由外到内、由显到隐逐层深入：")
    for s in [
        "第1轮：消灭显性漏洞（注释泄露、明文密码、Debug模式等直观问题）",
        "第2轮：修复逻辑漏洞（会话固定、计时攻击、限速粒度等架构问题）",
        "第3轮：合规与纵深防御（OWASP全覆盖、WCAG无障碍、审计日志）",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "3.2 第一轮修复（高危漏洞清零）", 2)
    add_body(doc, "目标：修复所有可以从页面直接观察到的安全缺陷。")
    add_table_with_style(doc,
        ["步骤", "漏洞", "设计思路", "实现方案"],
        [["1.1", "V-04 HTML注释泄露密码", "不信任任何静态文件", "删除login.html中的调试注释行"],
         ["1.2", "V-01 密码明文存储", "单向哈希保护凭据", "generate_password_hash()替代明文"],
         ["1.3", "V-05 密码明文展示", "最小化信息暴露", "sanitize_user_info()过滤password字段"],
         ["1.4", "V-03 Debug模式关闭", "生产环境禁用调试器", "debug=False"],
         ["1.5", "V-02 弱密钥", "不可预测的密钥生成", "os.urandom(24).hex()随机生成"],
         ["1.6", "V-06 密码明文比对", "引入哈希校验库", "check_password_hash()替代=="],
         ["1.7", "V-07 CSRF防护", "同步器令牌模式", "secrets.token_hex(32)生成token"],
         ["1.8", "V-08 暴力破解", "速率限制", "按IP记录，5次失败锁定5分钟"],
         ["1.9", "V-10 Cookie安全标记", "增强cookie防护", "HttpOnly=True, SameSite=Lax"],
         ["1.10", "V-14 输入校验", "前后端双重校验", "HTML maxlength + 服务端限制"],
         ["1.11", "V-16/V-17 安全头缺失", "纵深防御", "X-Content-Options + X-Frame-Options"]])
    add_body(doc, "验证结果：11个漏洞全部修复，首轮验证通过。", bold=True)

    add_heading_styled(doc, "3.3 第二轮修复（架构级漏洞加固）", 2)
    add_body(doc, "目标：修复第一轮未能覆盖的逻辑型漏洞及引入的回归问题。")
    add_table_with_style(doc,
        ["步骤", "漏洞", "设计思路", "实现方案"],
        [["2.1", "V-13 用户名枚举", "消除响应时间差异", "引入_DUMMY_HASH虚拟哈希；修复and短路"],
         ["2.2", "V-09 会话固定攻击", "登录时刷新会话", "session.clear()后重设username和token"],
         ["2.3", "V-08 暴力破解粒度", "精确到用户名", "限速字典由IP级改为用户名级"],
         ["2.4", "V-12 session过期时间", "带过期机制的会话", "PERMANENT_SESSION_LIFETIME=1800"],
         ["2.5", "V-11 Server版本隐藏", "消除协议层指纹", "覆盖WSGIRequestHandler.version_string"],
         ["2.6", "V-15 CSP优化", "升级为HTTP头", "移入after_request响应头并增强指令"]])
    add_body(doc, "验证结果：8个漏洞全部修复，第二轮验证通过。", bold=True)

    add_heading_styled(doc, "3.4 第三轮修复（合规审计与纵深防御）", 2)
    add_body(doc, "目标：满足 OWASP Top 10 2024 全部适配项，同时达到 WCAG 2.1 AA 无障碍标准。")
    add_table_with_style(doc,
        ["步骤", "标准", "设计思路", "实现方案"],
        [["3.1", "OWASP A1", "缺省拒绝的鉴权模式", "login_required装饰器 + 401/403错误页"],
         ["3.2", "OWASP A5", "安全头全面覆盖", "Referrer-Policy, Permissions-Policy, CSP增强"],
         ["3.3", "OWASP A9", "不可否认性", "审计日志记录登录成功/失败/锁定/CSRF异常"],
         ["3.4", "OWASP A7", "会话增强", "记录_last_active时间戳"],
         ["3.5", "WCAG 2.4.1", "键盘导航跳过链接", ".skip-link直达#main-content"],
         ["3.6", "WCAG 4.1.2", "ARIA Landmark标注", "role=navigation, aria-label, role=alert"],
         ["3.7", "WCAG 1.4.1", "高对比度焦点指示器", "outline:3px solid #ffbf47"],
         ["3.8", "WCAG 3.3.1", "错误实时播报", "aria-live=assertive aria-atomic=true"],
         ["3.9", "WCAG 1.4.3", "颜色对比度达标", "主色#1a1a1a，全部满足4.5:1对比度"]])
    add_body(doc, "验证结果：新增9项安全+无障碍措施，合规验收通过。", bold=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 4. 整改前后代码对比
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "4. 整改前后代码对比", 1)

    # 4.1
    add_heading_styled(doc, "4.1 密码存储与校验", 2)
    add_table_with_style(doc,
        ["对比项", "修复前（含V-01、V-06高危漏洞）", "修复后"],
        [["存储方式", "admin123 明文存储", "generate_password_hash() bcrypt哈希"],
         ["比对方式", "== 字符串直接比较", "check_password_hash() 常量时间比对"],
         ["返回前端", "完整用户对象含密码字段", "sanitize_user_info() 过滤密码"]])
    add_code_block(doc, '''USERS = {
    "admin": {
        "password": "admin123",          # 明文存储
    },
}
# 登录验证
if USERS[username]["password"] == password:  # 直接==比较''', "修复前代码")
    add_code_block(doc, '''USERS = {
    "admin": {
        "password": generate_password_hash("admin123"),  # bcrypt哈希
    },
}
# 登录验证
if check_password_hash(pwd_hash, password) and user_record:  # 恒等比对''', "修复后代码")

    # 4.2
    add_heading_styled(doc, "4.2 HTML注释与敏感信息", 2)
    add_table_with_style(doc,
        ["对比项", "修复前（含V-04高危漏洞）", "修复后"],
        [["HTML注释", "包含管理员账号密码", "整行删除，不留凭据痕迹"]])
    add_code_block(doc, '<!-- 调试信息 - 默认管理员账号 用户名: admin 密码: admin123 -->', "修复前代码")
    add_code_block(doc, '（该行已删除，生产环境不留任何凭据痕迹）', "修复后代码")

    # 4.3
    add_heading_styled(doc, "4.3 Session与认证安全", 2)
    add_table_with_style(doc,
        ["对比项", "修复前", "修复后"],
        [["secret_key", "dev-key-2025 硬编码", "os.urandom(24).hex() 随机生成"],
         ["cookie标记", "无安全配置", "HttpOnly, SameSite=Lax, 30分钟过期"],
         ["登录刷新", "直接写入用户名", "先清空旧会话再生成新会话"],
         ["登出清理", "仅删除个别键", "session.clear() 全量清空"]])
    add_code_block(doc, '''app.secret_key = "dev-key-2025"     # 硬编码弱密钥
session["username"] = username      # 未清空旧session
session.pop("username", None)       # 仅删除个别键''', "修复前代码")
    add_code_block(doc, '''app.secret_key = os.urandom(24).hex()
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    PERMANENT_SESSION_LIFETIME=1800,
)
session.clear()
session["username"] = username
session["_csrf_token"] = secrets.token_hex(32)
session["_last_active"] = time.time()
# 登出
session.clear()''', "修复后代码")

    # 4.4
    add_heading_styled(doc, "4.4 暴力破解与限速", 2)
    add_table_with_style(doc,
        ["对比项", "修复前", "修复后"],
        [["限速逻辑", "无限次尝试，无记录", "按用户名限速，跨IP防绕"],
         ["锁定阈值", "无", "5次连续失败触发"],
         ["锁定时长", "无", "300秒自动解锁"],
         ["内存保护", "无", "_prune_login_attempts()限制字典大小"]])
    add_code_block(doc, '''if username in USERS and USERS[username]["password"] == password:
    # 登录成功
else:
    # 登录失败，无记录''', "修复前代码")
    add_code_block(doc, '''LOGIN_ATTEMPTS = {}
MAX_ATTEMPTS = 5
LOCKOUT_TIME = 300
if username in LOGIN_ATTEMPTS:
    record = LOGIN_ATTEMPTS[username]
    if record["count"] >= MAX_ATTEMPTS:
        if time.time() - record["first_fail"] < LOCKOUT_TIME:
            return render_template("login.html", error="账号已被锁定")
        del LOGIN_ATTEMPTS[username]
LOGIN_ATTEMPTS[username] = {"count": count+1, "first_fail": now}''', "修复后代码")

    # 4.5
    add_heading_styled(doc, "4.5 计时攻击防御", 2)
    add_table_with_style(doc,
        ["对比项", "修复前（含V-13中危漏洞）", "修复后"],
        [["短路逻辑", "user_record and check_hash() 短路径", "check_hash() and user_record 始终执行"],
         ["虚拟哈希", "不存在用户直接返回", "用_DUMMY_HASH走完整bcrypt比对"]])
    add_code_block(doc, '''# 短路求值导致不存在的用户不执行哈希比对
if user_record and check_password_hash(password_hash, password):''', "修复前代码（含BUG）")
    add_code_block(doc, '''_DUMMY_HASH = generate_password_hash("__dummy__")
password_hash = user_record["password"] if user_record else _DUMMY_HASH
# 先执行哈希比对，再判断用户是否存在
if check_password_hash(password_hash, password) and user_record:''', "修复后代码")

    # 4.6
    add_heading_styled(doc, "4.6 HTTP安全响应头", 2)
    add_table_with_style(doc,
        ["安全头", "修复前", "修复后", "安全作用"],
        [["Server", "Werkzeug/3.1.8...", "WebServer", "隐藏版本指纹"],
         ["Content-Security-Policy", "无", "default-src 'self';...", "防XSS和数据注入"],
         ["X-Content-Type-Options", "无", "nosniff", "防MIME嗅探"],
         ["X-Frame-Options", "无", "DENY", "防点击劫持"],
         ["Referrer-Policy", "无", "strict-origin...", "控制Referer泄露"],
         ["Permissions-Policy", "无", "camera=(),...", "限制浏览器API权限"]])
    add_code_block(doc, '''# 无任何安全响应头
Server: Werkzeug/3.1.8 Python/3.13.12''', "修复前")
    add_code_block(doc, '''Content-Security-Policy: default-src 'self'; style-src 'self' 'unsafe-inline';
                        script-src 'self'; frame-ancestors 'none'; form-action 'self'; base-uri 'self'
X-Content-Type-Options: nosniff
X-Frame-Options: DENY
Referrer-Policy: strict-origin-when-cross-origin
Permissions-Policy: camera=(), microphone=(), geolocation=(), payment=()
Server: WebServer''', "修复后")

    # 4.7
    add_heading_styled(doc, "4.7 无障碍访问（WCAG 2.1）", 2)
    add_table_with_style(doc,
        ["无障碍特性", "适用标准", "受益群体", "作用说明"],
        [[".skip-link 跳过导航", "WCAG 2.4.1", "键盘/视障用户", "直接跳到主内容区"],
         ["role=navigation+aria-label", "WCAG 4.1.2", "屏幕阅读器用户", "标识导航区域名称"],
         ["role=alert+aria-live", "WCAG 3.3.1", "屏幕阅读器用户", "即时播报错误信息"],
         ["aria-required=true", "WCAG 3.3.2", "屏幕阅读器用户", "标识必填字段"],
         ["autocomplete", "WCAG 1.3.5", "认知障碍用户", "浏览器自动填充凭据"],
         ["outline:3px solid #ffbf47", "WCAG 2.4.7", "键盘用户", "高可见度焦点环"],
         ["prefers-reduced-motion", "WCAG 2.3.3", "前庭障碍用户", "关闭动效"]])
    add_code_block(doc, '''<body>
    <nav class="navbar">
        <div class="nav-brand">用户管理系统</div>
    </nav>
    <main class="container">
        {% block content %}{% endblock %}
    </main>
</body>''', "修复前代码（无无障碍属性）")
    add_code_block(doc, '''<body>
    <a href="#main-content" class="skip-link">跳到主内容</a>
    <nav role="navigation" aria-label="主导航">
        <div class="nav-brand">用户管理系统</div>
    </nav>
    <main id="main-content" role="main" aria-label="页面主要内容">
        <div role="alert" aria-live="assertive" aria-atomic="true">错误提示</div>
        {% block content %}{% endblock %}
    </main>
    <footer role="contentinfo" aria-label="页面底部信息">...</footer>
</body>''', "修复后代码（含无障碍属性）")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 5. 最终整改结果验证
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "5. 最终整改结果验证", 1)

    add_heading_styled(doc, "5.1 安全能力验证矩阵", 2)
    add_table_with_style(doc,
        ["测试项", "预期结果", "结果", "测试方法"],
        [["页面源码查看", "无账号密码暴露", "通过", "浏览器查看源代码"],
         ["密码哈希", "数据库不可逆", "通过", "检查存储的哈希值"],
         ["CSRF防护", "无token请求被拒", "通过", "curl提交缺token请求"],
         ["暴力破解", "5次失败后锁定", "通过", "连续6次错误密码"],
         ["Session固定", "登录前后session变更", "通过", "记录cookie前后对比"],
         ["计时攻击", "响应时间一致", "通过", "统计多次请求耗时"],
         ["Logout保护", "GET请求返回405", "通过", "curl发送GET /logout"],
         ["安全响应头", "全部6项存在", "通过", "curl检查响应头"],
         ["404错误页", "自定义404页", "通过", "访问不存在路径"],
         ["审计日志", "登录行为被记录", "通过", "查看flask.log"],
         ["焦点指示器", "Tab导航可见", "通过", "键盘Tab遍历"],
         ["屏幕阅读器", "ARIA属性正确", "通过", "WAVE/axe DevTools检测"]])

    add_heading_styled(doc, "5.2 修复后页面功能说明", 2)
    add_table_with_style(doc,
        ["页面", "功能说明", "安全状态"],
        [["首页（/）", "已登录显示用户信息；未登录显示请先登录", "密码不展示，401兜底"],
         ["登录页（/login）", "用户名+密码登录，CSRF token保护", "限速5次/5分钟"],
         ["登录流程", "session 30分钟过期，登录刷新会话", "审计日志全量记录"],
         ["登出流程", "POST提交带CSRF token", "GET请求被拒(405)"],
         ["错误页", "401/403/404/405/500统一模板", "不泄露服务器详情"]])

    add_heading_styled(doc, "5.3 修复效果总览", 2)
    add_table_with_style(doc,
        ["指标", "初始状态", "修复后状态"],
        [["漏洞总数", "19个（5高/8中/6低）", "0个已知漏洞"],
         ["安全响应头", "0个", "6项全部到位"],
         ["ARIA无障碍属性", "0个", "15+项"],
         ["审计日志", "未开启", "全量记录登录事件"],
         ["Server安全头", "暴露版本号", "已隐藏"]])

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 6. 总结与后续建议
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "6. 总结与后续建议", 1)

    add_heading_styled(doc, "6.1 修复总结", 2)
    add_body(doc, "本次审计对用户管理系统未登录页面进行了三轮递进式安全整改：")

    for title, desc in [
        ("1. 消灭显性漏洞（第1轮）",
         "删除了HTML中的明文凭据，将密码存储从明文切换到bcrypt哈希，关闭了Debug模式，引入了CSRF token保护机制和基础的暴力破解限速。"),
        ("2. 加固逻辑缺陷（第2轮）",
         "修复了会话固定攻击、计时侧信道用户名枚举、限速粒度过粗（从IP级优化到用户名级）、Server版本指纹泄露等架构层面的安全问题。"),
        ("3. 合规与纵深（第3轮）",
         "全面覆盖OWASP Top 10 2024的适配项，新增Referrer-Policy、Permissions-Policy等安全头，部署审计日志系统，同时使前端达到WCAG 2.1 AA无障碍标准。"),
    ]:
        add_body(doc, title, bold=True)
        add_body(doc, desc)

    add_body(doc, "最终通过全部验证项，已知漏洞清零。", bold=True)

    add_heading_styled(doc, "6.2 后续优化建议", 2)

    for rec in [
        ("建议一：升级为HTTPS协议",
         "当前系统仍运行于HTTP协议，session cookie和登录凭证以明文传输。",
         ["申请可信SSL证书（Let's Encrypt免费证书或商业证书）",
          "将 app.config 中 SESSION_COOKIE_SECURE 设为 True",
          "启用 Strict-Transport-Security 响应头（推荐 max-age=63072000）",
          "在Nginx/Caddy反向代理层终结TLS"]),
        ("建议二：引入数据库持久化存储",
         "当前用户数据存储在内存字典中，存在进程重启后数据丢失、无法水平扩展等限制。",
         ["引入 SQLite（轻量）或 PostgreSQL（生产）作为持久化存储",
          "用户密码哈希、角色信息存储于数据库",
          "登录失败记录写入持久化缓存（Redis），跨进程共享限速状态",
          "审计日志写入独立日志文件或数据库表，便于事后排查"]),
        ("建议三：增强用户认证机制",
         "当前仅依赖用户名+密码的单因素认证，存在口令泄露即账户失陷等风险。",
         ["引入密码强度检测（长度>=12、含大小写字母+数字+特殊字符）",
          "支持多因素认证（MFA）—— TOTP/短信验证码",
          "定期密码过期策略（90天强制更新）",
          "异常登录检测（异地IP、非常用设备、非工作时间登录告警）",
          "账户锁定后提供解锁流程（邮件验证码自助解锁）"]),
    ]:
        add_body(doc, rec[0], bold=True)
        add_body(doc, rec[1])
        for item in rec[2]:
            add_bullet(doc, item)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 附录
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "附录A：文件清单", 1)
    add_table_with_style(doc,
        ["文件", "路径", "说明"],
        [["主应用", "/workspace/user_management/app.py", "Flask应用入口，安全配置与路由"],
         ["基础模板", "/workspace/user_management/templates/base.html", "导航栏、安全头、无障碍骨架"],
         ["登录页", "/workspace/user_management/templates/login.html", "CSRF表单、ARIA无障碍"],
         ["首页", "/workspace/user_management/templates/index.html", "用户信息展示，登录态判断"],
         ["样式表", "/workspace/user_management/static/css/style.css", "高对比度焦点、响应式、减少动效"]])

    add_heading_styled(doc, "附录B：修复数据汇总", 1)
    add_table_with_style(doc,
        ["指标", "数值"],
        [["发现漏洞总数", "19个"],
         ["高危漏洞", "5个"],
         ["中危漏洞", "8个"],
         ["低危漏洞", "6个"],
         ["修复率", "100%"],
         ["新增安全响应头", "6个"],
         ["新增ARIA无障碍属性", "15+处"],
         ["新增代码行数（净增）", "约120行"]])

    # ── 结束 ──────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 报告结束 ——")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本报告由安全审计组自动生成，基于对用户管理系统 192.168.72.129:5000 的全量安全审计数据汇编。")
    run.font.size = Pt(8.5)
    run.font.name = '微软雅黑'
    run.font.color.rgb = C_MUTED

    out = "/workspace/user_management/阶段一/阶段一-全量安全审计报告.docx"
    doc.save(out)
    import os as _os
    sz = _os.path.getsize(out) / 1024
    print(f"✅ 报告已生成：{out}")
    print(f"   文件大小：{sz:.1f} KB")


if __name__ == "__main__":
    main()

# 18345434
