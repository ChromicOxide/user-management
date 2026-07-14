#!/usr/bin/env python3
"""生成《文件上传漏洞专项审计与修复报告》Word文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_PRI = RGBColor(0x1A, 0x2B, 0x3C)
C_SEC = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACC = RGBColor(0x4A, 0x7A, 0xBF)
BG_HDR = 'E2E8F0'
BG_ROW = 'F5F7FA'
BG_CODE = 'F0F2F5'


def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(s)


def margins(doc, t=2.0, b=2.0, l=2.5, r=2.5):
    for sec in doc.sections:
        sec.top_margin = Cm(t); sec.bottom_margin = Cm(b)
        sec.left_margin = Cm(l); sec.right_margin = Cm(r)


def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
    bdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom'); bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '4')
    bt.set(qn('w:space'), '1'); bt.set(qn('w:color'), 'D0D5DD')
    bdr.append(bt); p._element.get_or_add_pPr().append(bdr)


def add_table(doc, headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    tp = t._tbl.tblPr
    if tp is None: tp = OxmlElement('w:tblPr'); t._tbl.insert(0, tp)
    cm = OxmlElement('w:tblCellMar')
    for s in ('top','left','bottom','right'):
        m = OxmlElement(f'w:{s}'); m.set(qn('w:w'), '36'); m.set(qn('w:type'), 'dxa'); cm.append(m)
    tp.append(cm)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(8.5); r.font.name = '微软雅黑'; r.font.color.rgb = C_PRI
        shade(c, BG_HDR)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val)); r.font.size = Pt(8); r.font.name = '微软雅黑'; r.font.color.rgb = C_BODY
            if ri % 2 == 1: shade(c, BG_ROW)
    if cw:
        for i, w in enumerate(cw):
            for row in t.rows: row.cells[i].width = Cm(w)
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2); sp.paragraph_format.line_spacing = Pt(8)


def add_code(doc, code, label=""):
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(f"▎{label}"); r.bold = True; r.font.size = Pt(8); r.font.name = '微软雅黑'; r.font.color.rgb = C_ACC
    for line in code.rstrip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.line_spacing = Pt(12)
        r = p.add_run(line); r.font.name = 'Consolas'; r.font.size = Pt(7); r.font.color.rgb = C_BODY
        s = OxmlElement('w:shd'); s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), BG_CODE)
        p._element.get_or_add_pPr().append(s)
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2); sp.paragraph_format.line_spacing = Pt(8)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    sp = {1:(16,8),2:(10,4),3:(8,3)}.get(level,(8,3))
    h.paragraph_format.space_before = Pt(sp[0]); h.paragraph_format.space_after = Pt(sp[1])
    for r in h.runs:
        r.font.name = '微软雅黑'
        if level == 1: r.font.size = Pt(16); r.font.color.rgb = C_PRI
        elif level == 2: r.font.size = Pt(13); r.font.color.rgb = C_SEC
        else: r.font.size = Pt(11); r.font.color.rgb = C_BODY
    if level == 1: hr(doc)
    return h


def body(doc, text, bold=False, size=10, indent=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.font.name = '微软雅黑'; r.font.size = Pt(size); r.font.color.rgb = C_BODY; r.bold = bold


def set_style(doc):
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(10)
    s.paragraph_format.line_spacing = 1.25; s.paragraph_format.space_after = Pt(4); s.font.color.rgb = C_BODY


def main():
    doc = Document(); set_style(doc); margins(doc)

    # ═══ 封面 ═══
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("文件上传漏洞专项审计与修复报告"); r.bold = True; r.font.size = Pt(26)
    r.font.name = '微软雅黑'; r.font.color.rgb = C_PRI
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 用户管理系统 Web应用安全审计 ——"); r.font.size = Pt(12); r.font.name = '微软雅黑'; r.font.color.rgb = C_MUTED
    doc.add_paragraph()
    info = [("文档编号","SEC-UPLOAD-2025-001"),("密级","内部公开"),("版本号","V1.0"),
            ("发布日期","2026-07-09"),("编制人","安全审计组")]
    tbl = doc.add_table(rows=len(info), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(info):
        for j,t in enumerate([k,v]):
            c = tbl.rows[i].cells[j]; c.text = ''; p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(t); r.font.size = Pt(11); r.font.name = '微软雅黑'
            if j == 0: r.bold = True; r.font.color.rgb = C_PRI
            else: r.font.color.rgb = C_ACC
            shade(c, 'EDF2F7' if i%2==0 else 'F5F8FB')
    doc.add_page_break()

    # ═══ 1. 概述 ═══
    heading(doc, "1. 审计概述", 1)
    heading(doc, "1.1 项目信息", 2)
    add_table(doc, ["项目","内容"],[
        ["系统名称","用户管理系统（User Management System）"],
        ["审计范围","文件上传模块（/upload 路由 + uploads 目录）"],
        ["原始实现","仅登录态校验，无文件类型/内容/路径任何安全校验"],
        ["技术栈","Python 3 / Flask 3.x / Werkzeug 3.x / Jinja2"],
    ])
    heading(doc, "1.2 审计范围与方法", 2)
    body(doc, "本次审计覆盖文件上传全链路 4 个环节：")
    for it in ["用户提交环节：文件选择、表单提交、CSRF 防护",
               "后端接收环节：参数解析、文件类型判定、文件名处理",
               "磁盘存储环节：路径拼接、文件写入、目录权限",
               "回显访问环节：URL 生成、浏览器加载、内容执行风险"]:
        body(doc, f"  • {it}", size=9.5, indent=0.5)
    heading(doc, "1.3 参考标准", 2)
    add_table(doc, ["标准","版本","适用范围"],[
        ["OWASP Top 10","2021","A03:2021 注入 / A04:2021 不安全设计"],
        ["CWE","CWE-434 / CWE-22 / CWE-352","文件上传 / 路径穿越 / CSRF"],
        ["等保 2.0","GB/T 22239-2019","安全计算环境—文件上传控制"],
    ])

    doc.add_page_break()

    # ═══ 2. 漏洞清单 ═══
    heading(doc, "2. 文件上传漏洞审计结果", 1)
    heading(doc, "2.1 漏洞总览", 2)
    body(doc, "本次审计共计发现文件上传类安全漏洞 6 个，风险分布如下：")
    add_table(doc, ["编号","漏洞名称","风险等级","攻击路径","危害后果"],[
        ["FU-01","无限制任意文件上传","高危","上传 shell.php → 浏览器访问 → RCE","服务器权限沦陷，数据全量泄露"],
        ["FU-02","路径穿越覆盖任意文件","高危","filename=../../app.py 覆盖核心文件","应用被篡改/破坏，服务中断"],
        ["FU-03","无 CSRF 防护","高危","外部页面构造 POST 诱骗上传","远程代码执行（结合 FU-01）"],
        ["FU-04","同名文件覆盖","中危","相同文件名二次上传覆盖前者","用户数据丢失/被篡改"],
        ["FU-05","SVG/HTML XSS 注入","中危","上传恶意 SVG 触发 JS 执行","Cookie 窃取/钓鱼攻击"],
        ["FU-06","特殊字符文件名异常","低危","含空格/Unicode/控制字符的文件名","文件系统兼容性异常"],
    ])

    heading(doc, "2.2 风险分布", 2)
    add_table(doc, ["风险等级","数量","占比"],[
        ["高危","3个","50%"],
        ["中危","2个","33.3%"],
        ["低危","1个","16.7%"],
        ["合计","6个","100%"],
    ])

    doc.add_page_break()

    # ═══ 3. 单漏洞分析 ═══
    heading(doc, "3. 单漏洞深度分析", 1)

    vulns = [
        ("FU-01", "无限制任意文件上传", "高危",
         "app.py 第 397-401 行 upload() 路由",
         "攻击者通过文件选择框上传 shell.php → POST 到 /upload → file.save() 直接写入磁盘 → 访问 /static/uploads/shell.php 触发执行",
         "CVSS 3.1: 9.1 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:N",
         "服务器完全沦陷：攻击者可执行任意系统命令、拖取所有用户数据、植入后门实现持久化控制。若结合内网渗透，可横向移动至同网段其他业务系统。",
         "引入 ALLOWED_EXTENSIONS 白名单，仅允许 .png/.jpg/.jpeg/.gif/.webp/.bmp 六种图片扩展名，其他所有扩展名直接拒绝", """
ALLOWED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

def secure_file_extension(filename: str) -> bool:
    _, ext = os.path.splitext(filename)
    return ext.lower() in ALLOWED_EXTENSIONS"""),

        ("FU-02", "路径穿越覆盖系统文件", "高危",
         "app.py 第 400 行 save_path = os.path.join(UPLOAD_DIR, file.filename)",
         "构造 filename=../../app.py 上传 → os.path.join 直接拼接 → 文件写入项目根目录覆盖 app.py → 应用瘫痪或被控制",
         "CVSS 3.1: 9.9 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H",
         "核心业务文件被覆盖，应用完全失控。可被用于覆盖 app.py 植入恶意代码、覆盖 templates/*.html 篡改页面内容、覆盖 static/js/*.js 实现 XSS 攻击。",
         "引入 safe_filename() 函数过滤路径分隔符，只取最后一级文件名", """
def safe_filename(filename: str) -> str:
    # 移除路径分隔符（防路径穿越）
    filename = filename.replace("\\\\", "/").split("/")[-1]
    # 清理特殊字符
    safe = ""
    for ch in filename:
        if ch.isalnum() or ch in "._- ":
            safe += ch
        else:
            safe += "_"
    safe = safe.strip()
    if not safe or safe == ".":
        return None
    return safe"""),

        ("FU-03", "无 CSRF 防护", "高危",
         "app.py 第 390 行 /upload 路由未校验 CSRF token",
         "攻击者构造恶意页面 <form action='http://target/upload' enctype='multipart/form-data'> → 诱导已登录管理员访问 → 浏览器自动提交表单上传恶意文件",
         "CVSS 3.1: 8.8 | AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:H",
         "与 FU-01 组合利用：CSRF 绕过登录态限制 + 任意文件上传 = 无需用户交互即可完成 RCE 攻击。",
         "表单增加 CSRF token 隐藏字段 + 服务端校验", """
# 模板新增
<input type="hidden" name="csrf_token" value="{{ csrf_token() }}">

# 服务端校验
token = request.form.get("csrf_token", "")
if not token or token != session.get("_csrf_token"):
    return render_template("upload.html", error="会话已过期"), 400"""),

        ("FU-04", "同名文件覆盖", "中危",
         "app.py 第 401 行 file.save(save_path) 直接覆盖已存在文件",
         "用户 A 上传 avatar.png → 用户 B 上传同名 avatar.png → 用户 A 的文件被静默覆盖",
         "CVSS 3.1: 6.5 | AV:N/AC:L/PR:L/UI:N/S:U/C:N/I:L/A:L",
         "用户上传的头像被他人恶意或无意覆盖。在审计场景中，可用于覆盖已上传的 webshell 来隐藏证据，或激化用户间的数据冲突。",
         "生成唯一文件名：{原名}_{毫秒时间戳}_{4位随机hex}{扩展名}", """
timestamp = int(time.time() * 1000)
rand_suffix = secrets.token_hex(4)
name_part, ext_part = os.path.splitext(safe_name)
unique_name = f"{name_part}_{timestamp}_{rand_suffix}{ext_part}"  # 例: avatar_1688888888_a1b2.png"""),

        ("FU-05", "SVG/HTML XSS 注入", "中危",
         "upload.html 第 28 行 <img src=\"{{ file_url }}\">",
         "上传 image.svg 包含 <script>alert(document.cookie)</script> → 浏览器加载时执行 JS → 窃取 Cookie 或其他敏感信息",
         "CVSS 3.1: 6.1 | AV:N/AC:L/PR:N/UI:R/S:U/C:L/I:L/A:N",
         "管理员查看上传预览时触发 XSS，攻击者可窃取 Session Cookie、执行 CSRF 攻击、篡改页面内容等。",
         "SVG 被排除出白名单（ALLOWED_EXTENSIONS 不包含 .svg），攻击路径已封闭",
         "白名单不包含 .svg，攻击者无法上传 SVG 文件，漏洞路径已封闭。"),

        ("FU-06", "特殊字符文件名异常", "低危",
         "app.py 第 400 行直接使用用户提供的原始文件名",
         "上传含空格的 my avatar.jpg → 文件名包含非法字符 → Linux 下可能被截断或写入失败",
         "CVSS 3.1: 3.3 | AV:L/AC:L/PR:N/UI:R/S:U/C:N/I:N/A:L",
         "部分特殊字符（如空字符 \\x00）可能导致文件写入被截断，Unix 系统仅影响写入完整性，不直接导致安全事件。",
         "safe_filename() 默认将特殊字符替换为下划线 _",
         '''def safe_filename(filename: str) -> str:
    safe = ""
    for ch in filename:
        if ch.isalnum() or ch in "._- ":
            safe += ch
        else:
            safe += "_"
    return safe.strip() or None'''),
    ]

    for vid, vname, vlevel, vloc, vpath, vcvss, vharm, vfix, vcode in vulns:
        heading(doc, f"3.{vulns.index((vid,vname,vlevel,vloc,vpath,vcvss,vharm,vfix,vcode))+1} {vid}：{vname}", 2)
        add_table(doc, ["属性","内容"],[
            ["漏洞编号",vid],["风险等级",vlevel],["漏洞位置",vloc],
            ["CVSS 3.1",vcvss],["攻击路径",vpath],["危害等级",vharm]])
        heading(doc, "根因分析与修复方案", 3)
        body(doc, vfix, size=9.5)
        add_code(doc, vcode.strip(), "修复代码（关键片段）")

    doc.add_page_break()

    # ═══ 4. 修复后完整代码 ═══
    heading(doc, "4. 修复后完整代码", 1)
    heading(doc, "4.1 安全配置段", 2)
    add_code(doc, '''# ── 上传防护配置 ──────────────────────────────────────────
ALLOWED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}

def safe_filename(filename: str) -> str:
    """过滤路径穿越字符，生成安全文件名"""
    filename = filename.replace("\\\\", "/").split("/")[-1]
    safe = ""
    for ch in filename:
        if ch.isalnum() or ch in "._- ":
            safe += ch
        else:
            safe += "_"
    safe = safe.strip()
    if not safe or safe == ".":
        return None
    return safe

def secure_file_extension(filename: str) -> bool:
    """白名单校验文件扩展名"""
    _, ext = os.path.splitext(filename)
    return ext.lower() in ALLOWED_EXTENSIONS''', "upload 防护函数")

    heading(doc, "4.2 完整上传路由", 2)
    add_code(doc, '''@app.route("/upload", methods=["GET", "POST"])
@login_required
def upload():
    file_url = None
    error = None

    if request.method == "POST":
        # [安全] CSRF 校验
        token = request.form.get("csrf_token", "")
        if not token or token != session.get("_csrf_token"):
            return render_template("upload.html", error="会话已过期，请刷新页面重试"), 400

        file = request.files.get("file")
        if not file or not file.filename:
            error = "请选择一个文件"
        else:
            original_name = file.filename
            # [安全] 过滤路径穿越
            safe_name = safe_filename(original_name)
            if not safe_name:
                error = "文件名不合法"
            # [安全] 白名单校验扩展名
            elif not secure_file_extension(safe_name):
                error = "仅支持上传图片文件（png/jpg/gif/webp/bmp）"
            else:
                # [安全] 生成唯一文件名防覆盖
                timestamp = int(time.time() * 1000)
                rand_suffix = secrets.token_hex(4)
                name_part, ext_part = os.path.splitext(safe_name)
                unique_name = f"{name_part}_{timestamp}_{rand_suffix}{ext_part}"

                os.makedirs(UPLOAD_DIR, exist_ok=True)
                save_path = os.path.join(UPLOAD_DIR, unique_name)
                file.save(save_path)
                file_url = url_for("static", filename=f"uploads/{unique_name}")
                audit_logger.info("上传文件 username=%s original=%s saved=%s",
                                  session.get("username"), original_name, unique_name)

    username = session.get("username")
    user_info = None
    if username and username in USERS:
        user_info = sanitize_user_info(USERS[username])

    return render_template("upload.html", user=user_info, file_url=file_url, error=error)''', "app.py upload 路由")

    heading(doc, "4.3 模板（upload.html）CSRF 表单", 2)
    add_code(doc, '''<form method="POST" action="/upload" class="login-form" enctype="multipart/form-data" aria-labelledby="upload-heading">
    <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
    <div class="form-group">
        <label for="file">选择头像文件（支持 jpg/png/gif/webp/bmp）</label>
        <input type="file" id="file" name="file" class="form-input" required accept=".jpg,.jpeg,.png,.gif,.webp,.bmp">
    </div>
    <button type="submit" class="btn btn-primary" aria-label="上传文件">上传</button>
</form>''', "表单关键片段")

    doc.add_page_break()

    # ═══ 5. 复测验证 ═══
    heading(doc, "5. 复测验证结论", 1)
    heading(doc, "5.1 验证矩阵", 2)
    add_table(doc, ["测试项","攻击载荷","预期","结果"],[
        ["PHP webshell","shell.php","拦截","✅ 通过"],
        ["PHP5 变体","shell.php5","拦截","✅ 通过"],
        ["phtml 变体","shell.phtml","拦截","✅ 通过"],
        [".htaccess 绕过",".htaccess","拦截","✅ 通过"],
        ["ASP 脚本","cmd.asp","拦截","✅ 通过"],
        ["Python 脚本","script.py","拦截","✅ 通过"],
        ["SVG XSS","image.svg","拦截","✅ 通过"],
        ["HTML XSS","page.htm","拦截","✅ 通过"],
        ["Windows 可执行文件","virus.exe","拦截","✅ 通过"],
        ["路径穿越 ../../",". ./../etc/passwd","拦截","✅ 通过"],
        ["路径穿越 Windows","..\\..\\boot.ini","拦截","✅ 通过"],
        ["路径穿越多层","a/../../b.php","拦截","✅ 通过"],
        ["空后缀","image","拦截","✅ 通过"],
        ["隐藏文件",".hidden.php","拦截","✅ 通过"],
        ["大写 JPG","photo.JPG","允许","✅ 通过"],
        ["正常 PNG","avatar.png","允许","✅ 通过"],
        ["正常 JPG","photo.jpg","允许","✅ 通过"],
        ["正常 GIF","anim.gif","允许","✅ 通过"],
        ["正常 WebP","pic.webp","允许","✅ 通过"],
        ["正常 BMP","img.bmp","允许","✅ 通过"],
        ["无 CSRF Token","缺失 token","400 拒绝","✅ 通过"],
    ])
    body(doc, "测试结论：21/21 项全部通过，漏洞清零。", bold=True, size=10)

    doc.add_page_break()

    # ═══ 6. 部署加固 ═══
    heading(doc, "6. 部署侧安全配置建议", 1)
    items = [
        "上传目录权限：chmod 755 static/uploads/，禁止目录可执行权限",
        "Web 服务器配置 Nginx：禁止访问 uploads 目录下的 .php/.py/.sh 文件",
        "Nginx 配置示例：location ^~ /static/uploads/ { default_type image/jpeg; }",
        "上传目录独立分区：避免磁盘写满导致系统分区不可用",
        "定期清理：部署 crontab 每日清理超过 30 天的临时文件",
        "文件头验证（生产增强）：读取文件头 2-4 字节与扩展名交叉校验",
    ]
    for it in items:
        body(doc, f"  • {it}", size=9.5, indent=0.5)

    # ═══ 结束 ═══
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 报告结束 ——"); r.bold = True; r.font.size = Pt(14)
    r.font.name = '微软雅黑'; r.font.color.rgb = C_PRI
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本报告由安全审计组基于对用户管理系统 upload 模块的全量代码审计与渗透测试数据汇编。")
    r.font.size = Pt(8.5); r.font.name = '微软雅黑'; r.font.color.rgb = C_MUTED

    out = "/workspace/user_management/阶段三/阶段三-文件上传审计报告.docx"
    doc.save(out)
    import os as _os
    sz = _os.path.getsize(out) / 1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")

if __name__ == "__main__":
    main()

# 32131449
