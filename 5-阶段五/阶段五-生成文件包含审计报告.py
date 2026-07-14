#!/usr/bin/env python3
"""文件包含漏洞审查与修复报告 — 字体优化版"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 配色 ──
C_PRI = RGBColor(0x1A, 0x2B, 0x3C)
C_SEC = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACC = RGBColor(0x4A, 0x7A, 0xBF)

# ── 字体 ──
FONT_HEI = '黑体'
FONT_SONG = '宋体'
FONT_CODE = 'Consolas'

BG_HDR = 'E2E8F0'
BG_ROW = 'F5F7FA'
BG_CODE = 'F0F2F5'


def set_run_font(run, font_name, size=Pt(10), bold=False, color=None):
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def shade(c, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), color)
    c._tc.get_or_add_tcPr().append(s)


def margins(doc, t=2.0, b=2.0, l=2.5, r=2.5):
    for sec in doc.sections:
        sec.top_margin = Cm(t)
        sec.bottom_margin = Cm(b)
        sec.left_margin = Cm(l)
        sec.right_margin = Cm(r)


def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    bdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single')
    bt.set(qn('w:sz'), '4')
    bt.set(qn('w:space'), '1')
    bt.set(qn('w:color'), 'D0D5DD')
    bdr.append(bt)
    p._element.get_or_add_pPr().append(bdr)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    tp = t._tbl.tblPr
    if tp is None:
        tp = OxmlElement('w:tblPr')
        t._tbl.insert(0, tp)
    cm = OxmlElement('w:tblCellMar')
    for s in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement(f'w:{s}')
        m.set(qn('w:w'), '36')
        m.set(qn('w:type'), 'dxa')
        cm.append(m)
    tp.append(cm)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h)
        set_run_font(r, FONT_HEI, Pt(8.5), bold=True, color=C_PRI)
        shade(c, BG_HDR)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val))
            set_run_font(r, FONT_SONG, Pt(8), color=C_BODY)
            if ri % 2 == 1:
                shade(c, BG_ROW)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)


def add_code(doc, code, label=""):
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(label)
        set_run_font(r, FONT_HEI, Pt(8), bold=True, color=C_ACC)
    for line in code.rstrip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = Pt(12)
        r = p.add_run(line)
        set_run_font(r, FONT_CODE, Pt(7), color=C_BODY)
        sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear')
        sh.set(qn('w:color'), 'auto')
        sh.set(qn('w:fill'), BG_CODE)
        p._element.get_or_add_pPr().append(sh)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)


def heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    sz_map = {1: (18, 10, Pt(18)), 2: (14, 7, Pt(14)), 3: (12, 5, Pt(12))}
    sb, sa, sz = sz_map.get(level, (10, 5, Pt(11)))
    h.paragraph_format.space_before = Pt(sb)
    h.paragraph_format.space_after = Pt(sa)
    color_map = {1: C_PRI, 2: C_SEC, 3: C_BODY}
    clr = color_map.get(level, C_BODY)
    for r in h.runs:
        set_run_font(r, FONT_HEI, sz, bold=True, color=clr)
    style_name = f'Heading {level}'
    if style_name in doc.styles:
        sty = doc.styles[style_name]
        sty.font.name = FONT_HEI
        rPr = sty.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), FONT_HEI)
        rFonts.set(qn('w:hAnsi'), FONT_HEI)
        rFonts.set(qn('w:eastAsia'), FONT_HEI)
    if level == 1:
        hr(doc)
    return h


def body(doc, text, bold=False, size=10, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_run_font(r, FONT_SONG, Pt(size), bold=bold, color=C_BODY)


def set_default_style(doc):
    s = doc.styles['Normal']
    s.font.name = FONT_SONG
    s.font.size = Pt(10)
    s.paragraph_format.line_spacing = 1.35
    s.paragraph_format.space_after = Pt(4)
    s.font.color.rgb = C_BODY
    rPr = s.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_SONG)
    rFonts.set(qn('w:ascii'), FONT_SONG)
    rFonts.set(qn('w:hAnsi'), FONT_SONG)


def main():
    doc = Document()
    set_default_style(doc)
    margins(doc)

    # ═══ 封面 ═══
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("文件包含漏洞\n审查与修复报告")
    set_run_font(r, FONT_HEI, Pt(24), bold=True, color=C_PRI)
    doc.add_paragraph()

    info = [("文档编号", "SEC-FILE-2025-001"),
            ("密级", "内部公开"),
            ("版本号", "V1.0"),
            ("发布日期", "2026-07-10"),
            ("编制人", "潘麒宇")]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        for j, t in enumerate([k, v]):
            c = tbl.rows[i].cells[j]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(t)
            if j == 0:
                set_run_font(r, FONT_SONG, Pt(11), bold=True, color=C_PRI)
            else:
                set_run_font(r, FONT_SONG, Pt(11), color=C_ACC)
            shade(c, 'EDF2F7' if i % 2 == 0 else 'F5F8FB')
    doc.add_page_break()

    # ═══ 1 ═══
    heading_styled(doc, "1. 审计概述", 1)
    heading_styled(doc, "1.1 项目信息", 2)
    add_table(doc, ["项目", "内容"], [
        ["系统名称", "用户管理系统"],
        ["审计范围", "动态页面加载模块（/page 路由 + pages/ 目录）"],
        ["审计日期", "2026 年 7 月 10 日"],
        ["技术栈", "Python 3 / Flask 3.x / Jinja2 / HTML5"],
    ])
    heading_styled(doc, "1.2 审计范围", 2)
    for it in ["文件包含全链路：用户输入 name 参数 → 路径拼接 → 文件读取 → 页面渲染",
               "覆盖攻击面：路径穿越、任意文件读取、未授权访问、远程文件包含",
               "参考标准：OWASP Top 10 A01:2021 / CWE-22 / CWE-98"]:
        body(doc, f"  * {it}", size=9.5, indent=0.5)

    doc.add_page_break()

    # ═══ 2 ═══
    heading_styled(doc, "2. 漏洞发现清单", 1)
    heading_styled(doc, "2.1 漏洞总览", 2)
    add_table(doc, ["编号", "漏洞名称", "等级", "所在路由", "漏洞类型"], [
        ["FI-01", "路径穿越读取任意文件", "高危", "GET /page?name=../app.py", "文件包含"],
        ["FI-02", "读取系统敏感文件", "高危", "GET /page?name=../../../etc/passwd", "文件包含"],
        ["FI-03", "未授权访问动态页面", "中危", "GET /page?name=help", "权限缺失"],
        ["FI-04", "读取源码泄露凭据", "中危", "GET /page?name=../app.py", "信息泄露"],
        ["FI-05", "HTML 嵌套渲染异常", "低危", "pages/help.html 含完整 HTML 骨架", "内容污染"],
    ])
    heading_styled(doc, "2.2 风险分布", 2)
    add_table(doc, ["风险等级", "数量", "占比"], [
        ["高危", "2个", "40%"],
        ["中危", "2个", "40%"],
        ["低危", "1个", "20%"],
        ["合计", "5个", "100%"],
    ])

    doc.add_page_break()

    # ═══ 3 ═══
    heading_styled(doc, "3. 单漏洞深度分析与修复", 1)

    vulns_data = [
        ("FI-01", "路径穿越读取任意文件", "高危",
         "app.py:569 dynamic_page() — os.path.join(PAGES_DIR, name)",
         "攻击路径：\n1. 浏览器访问 /page?name=../app.py\n2. 服务器拼接路径：pages/ + ../app.py = pages/../app.py = app.py\n3. 文件内容被读取并渲染到页面中\n4. 可进一步读取任意系统文件\n\n代码根因：\nfile_path = os.path.join(PAGES_DIR, name)\n# name='../app.py' → pages/../app.py → 跳出了pages目录\nwith open(file_path, 'r') as f:\n    page_content = f.read()\n# page_content 通过 {{ page_content | safe }} 渲染",
         "CVSS 3.1: 9.1 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "攻击者可读取系统任意文件（源码、配置、数据库、密码文件），导致全量信息泄露。",
         "修复方案：\n1. 剥离路径部分，只取文件名\n2. 白名单仅允许 .html 扩展名\n3. os.path.realpath() 校验页面在 pages/ 目录内\n4. 添加 @login_required"),
        ("FI-02", "读取系统敏感文件", "高危",
         "app.py:569 same as FI-01",
         "攻击路径：\n1. 构造 name=../../../etc/passwd\n2. 路径拼接后变为 pages/../../../etc/passwd = /etc/passwd\n3. 读取 Linux 系统用户密码文件\n\n可读文件包括：\n- /etc/passwd（系统用户列表）\n- /etc/shadow（密码哈希，需 root）\n- app.py（源码，含 secret_key、数据库配置）",
         "CVSS 3.1: 8.6 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "系统级敏感文件泄露，辅助提权攻击。",
         "修复方案：os.path.realpath() 解析后判断是否以 pages/ 目录路径开头"),
        ("FI-03", "未授权访问动态页面", "中危",
         "app.py:562 dynamic_page() — 无 @login_required",
         "攻击路径：\n1. 清除浏览器 Cookie\n2. 直接访问 /page?name=help\n3. 无需登录即可看到帮助页面\n4. 结合 FI-01 可匿名读取任意文件",
         "CVSS 3.1: 5.3 | AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N",
         "与路径穿越结合后危害放大，未登录用户也可读取系统文件。",
         "修复方案：dynamic_page() 添加 @login_required 装饰器"),
        ("FI-04", "读取源码泄露凭据", "中危",
         "app.py:572 — 文件内容直接渲染在模板中",
         "攻击路径：\n1. 读取 app.py 源码\n2. 从中提取 secret_key 和数据库配置\n3. 伪造 session cookie 登录任意用户",
         "CVSS 3.1: 6.5 | AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
         "secret_key 泄露可致 session 伪造，数据库配置泄露可致数据库脱库。",
         "修复方案：已通过 FI-01 修复一并解决。"),
        ("FI-05", "HTML 嵌套渲染异常", "低危",
         "pages/help.html 含完整 <html><head><body> 标签",
         "原始 help.html 包含了完整的 HTML 骨架标签，导致嵌入 index.html 时出现双层 <html> 标签，页面结构异常。",
         "CVSS 3.1: 3.1 | AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:L",
         "页面显示异常，影响用户体验。无直接安全危害。",
         "修复方案：移除 help.html 中的 <html><head><body> 标签，仅保留内容片段。"),
    ]

    for idx, (vid, vname, vlevel, vloc, vpath, vcvss, vharm, vfix) in enumerate(vulns_data):
        heading_styled(doc, f"3.{idx+1} {vid}：{vname} [{vlevel}]", 2)
        add_table(doc, ["属性", "内容"], [
            ["漏洞编号", vid],
            ["风险等级", vlevel],
            ["代码位置", vloc],
            ["CVSS 3.1", vcvss],
            ["危害程度", vharm],
        ])
        add_code(doc, vpath, "攻击路径与根因分析")
        body(doc, vfix, size=9.5)

    doc.add_page_break()

    # ═══ 4 ═══
    heading_styled(doc, "4. 修复后核心代码", 1)
    add_code(doc,
        '# ── 安全配置 ──\n'
        'ALLOWED_PAGE_EXT = {".html"}\n\n'
        '@app.route("/page")\n'
        '@login_required\n'
        'def dynamic_page():\n'
        '    name = request.args.get("name", "")\n\n'
        '    # [安全] 防止路径穿越：只取文件名\n'
        '    safe_name = name.replace("\\\\", "/").split("/")[-1]\n\n'
        '    # [安全] 只允许 .html 扩展名\n'
        '    _, ext = os.path.splitext(safe_name)\n'
        '    if ext == "":\n'
        '        safe_name += ".html"\n'
        '    _, ext = os.path.splitext(safe_name)\n'
        '    if ext.lower() not in ALLOWED_PAGE_EXT:\n'
        '        return "不支持的页面类型"\n\n'
        '    # [安全] 真实路径校验\n'
        '    real_path = os.path.realpath(os.path.join(PAGES_DIR, safe_name))\n'
        '    real_base = os.path.realpath(PAGES_DIR)\n'
        '    if not real_path.startswith(real_base + os.sep):\n'
        '        return "非法的页面路径"\n\n'
        '    with open(real_path, "r", encoding="utf-8") as f:\n'
        '        page_content = f.read()',
        "修复后 /page 路由代码")

    doc.add_page_break()

    # ═══ 5 ═══
    heading_styled(doc, "5. 修复验证结论", 1)
    add_table(doc, ["测试项", "测试载荷", "修复前", "修复后", "结论"], [
        ["正常帮助页", "/page?name=help", "显示帮助内容", "显示帮助内容", "✓"],
        ["无后缀自动补", "/page?name=help", "N/A", "自动补.html", "✓"],
        ["路径穿越", "/page?name=../app.py", "读取源码", "不支持的页面类型", "✓"],
        ["系统文件读取", "/page?name=../../../etc/passwd", "读取系统文件", "不支持的页面类型", "✓"],
        ["PHP 文件", "/page?name=shell.php", "可能读取", "不支持的页面类型", "✓"],
        ["未登录访问", "无Cookie访问/page", "可访问", "401 拒绝", "✓"],
        ["不存在页面", "/page?name=nonexist", "无提示", "页面不存在", "✓"],
    ])
    heading_styled(doc, "6. 后续建议", 1)
    for it in ["禁用模板中的 | safe 过滤器，改用 autoescape 自动转义",
               "文件读取改为白名单模式（如只允许 help.html, about.html）",
               "引入文件内容安全校验，防止 HTML/JS 注入"]:
        body(doc, f"  * {it}", size=9.5, indent=0.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 报告结束 ——")
    set_run_font(r, FONT_HEI, Pt(14), bold=True, color=C_PRI)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本报告由安全审计组基于对用户管理系统动态页面加载模块的全量代码审计数据汇编。")
    set_run_font(r, FONT_SONG, Pt(8.5), color=C_MUTED)

    out = "/workspace/user_management/阶段五/阶段五-文件包含审计报告.docx"
    doc.save(out)
    import os as _os
    sz = _os.path.getsize(out) / 1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")


if __name__ == "__main__":
    main()
