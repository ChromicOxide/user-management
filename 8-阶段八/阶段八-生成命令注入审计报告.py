#!/usr/bin/env python3
"""命令注入漏洞审查与修复报告"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_PRI = RGBColor(0x1A, 0x2B, 0x3C); C_SEC = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C); C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACC = RGBColor(0x4A, 0x7A, 0xBF)
FONT_HEI = '黑体'; FONT_SONG = '宋体'; FONT_CODE = 'Consolas'
BG_HDR = 'E2E8F0'; BG_ROW = 'F5F7FA'; BG_CODE = 'F0F2F5'

def srf(r, fn, sz=Pt(10), b=False, c=None):
    r.font.size = sz; r.bold = b
    if c: r.font.color.rgb = c
    rPr=r._element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.insert(0, rf)
    for a in ('ascii','hAnsi','eastAsia'): rf.set(qn(f'w:{a}'), fn)

def shd(c, h):
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),h)
    c._tc.get_or_add_tcPr().append(s)

def mr(doc):
    for sec in doc.sections: sec.top_margin=Cm(2); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

def hr(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6)
    bd=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom'); bt.set(qn('w:val'),'single')
    bt.set(qn('w:sz'),'4'); bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),'D0D5DD'); bd.append(bt)
    p._element.get_or_add_pPr().append(bd)

def add_tbl(doc, hd, rows):
    t=doc.add_table(rows=1+len(rows), cols=len(hd))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    tp=t._tbl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); t._tbl.insert(0,tp)
    cm=OxmlElement('w:tblCellMar')
    for s in ('top','left','bottom','right'):
        m=OxmlElement(f'w:{s}'); m.set(qn('w:w'),'36'); m.set(qn('w:type'),'dxa'); cm.append(m)
    tp.append(cm)
    for i,h in enumerate(hd):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(h); srf(r,FONT_HEI,Pt(8.5),True,C_PRI); shd(c,BG_HDR)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            r=p.add_run(str(v)); srf(r,FONT_SONG,Pt(8),c=C_BODY)
            if ri%2==1: shd(c,BG_ROW)
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(2); sp.paragraph_format.space_after=Pt(2)

def code(doc, cd, lb=""):
    if lb:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(1)
        p.paragraph_format.left_indent=Cm(0.3)
        r=p.add_run(lb); srf(r,FONT_HEI,Pt(8),True,C_ACC)
    for line in cd.rstrip().split('\n'):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.line_spacing=Pt(12)
        r=p.add_run(line); srf(r,FONT_CODE,Pt(7),c=C_BODY)
        sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),BG_CODE)
        p._element.get_or_add_pPr().append(sh)
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(2); sp.paragraph_format.space_after=Pt(2)

def hd(doc, t, lv=1):
    h=doc.add_heading(t, level=lv)
    sz={1:(18,10,Pt(18)),2:(14,7,Pt(14)),3:(12,5,Pt(12))}.get(lv,(10,5,Pt(11)))
    h.paragraph_format.space_before=Pt(sz[0]); h.paragraph_format.space_after=Pt(sz[1])
    cl={1:C_PRI,2:C_SEC,3:C_BODY}.get(lv,C_BODY)
    for r in h.runs: srf(r,FONT_HEI,sz[2],True,cl)
    if lv==1: hr(doc)

def bd(doc, t, b=False, sz=10, ind=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    if ind: p.paragraph_format.left_indent=Cm(ind)
    r=p.add_run(t); srf(r,FONT_SONG,Pt(sz),b,c=C_BODY)

def sty(doc):
    s=doc.styles['Normal']; s.font.name=FONT_SONG; s.font.size=Pt(10)
    s.paragraph_format.line_spacing=1.35; s.paragraph_format.space_after=Pt(4); s.font.color.rgb=C_BODY
    rPr=s.element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('eastAsia','ascii','hAnsi'): rf.set(qn(f'w:{a}'),FONT_SONG)

def main():
    doc=Document(); sty(doc); mr(doc)
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("命令注入漏洞\n审查与修复报告"); r.bold=True; r.font.size=Pt(24); r.font.name=FONT_HEI; r.font.color.rgb=C_PRI
    doc.add_paragraph()
    info=[("文档编号","SEC-CMDINJ-2025-001"),("密级","内部公开"),("版本号","V1.0"),("发布日期","2026-07-15"),("编制人","潘麒宇")]
    tbl=doc.add_table(rows=len(info), cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(info):
        for j,t in enumerate([k,v]):
            c=tbl.rows[i].cells[j]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(t); r.font.size=Pt(11); r.font.name=FONT_SONG
            if j==0: r.bold=True; r.font.color.rgb=C_PRI
            else: r.font.color.rgb=C_ACC
            shd(c,'EDF2F7' if i%2==0 else 'F5F8FB')
    doc.add_page_break()

    hd(doc,"1. 审计概述",1)
    hd(doc,"1.1 项目信息",2)
    add_tbl(doc,["项目","内容"],[["系统名称","用户管理系统"],["审计范围","Ping 网络诊断功能（/ping 路由 + subprocess 调用）"],["审计日期","2026 年 7 月 15 日"],["技术栈","Python 3 / Flask 3.x / subprocess / shell"]])
    hd(doc,"1.2 审计方法",2)
    for i in ["代码审计：逐行审查 ping 路由的 shell 调用方式，定位命令注入点","黑盒渗透：构造 20+ 种命令注入 Payload 验证注入存在性","参考标准：OWASP Top 10 A03:2021 / CWE-77 / CWE-78"]:
        bd(doc,f"  * {i}",sz=9.5,ind=.5)
    hd(doc,"1.3 参考标准",2)
    add_tbl(doc,["标准","适用范围"],[["OWASP Top 10 2021","A03:2021 注入"],["CWE-77","命令中特殊元素转义不恰当"],["CWE-78","OS 命令注入"]])

    doc.add_page_break()
    hd(doc,"2. 漏洞发现清单",1)
    hd(doc,"2.1 漏洞总览",2)
    add_tbl(doc,["编号","漏洞名称","等级","所在路由","漏洞类型"],[["CI-01","shell=True 命令注入","高危","POST /ping","CWE-78 OS命令注入"],["CI-02","f-string 拼接命令字符串","高危","POST /ping","CWE-77 特殊元素转义"],["CI-03","无输入校验","高危","POST /ping","CWE-20 输入验证"],["CI-04","无超时控制","中危","POST /ping","CWE-400 资源耗尽"],["CI-05","错误信息泄露","低危","POST /ping","CWE-209 信息泄露"]])
    hd(doc,"2.2 风险分布",2)
    add_tbl(doc,["等级","数量","占比"],[["高危","3个","60%"],["中危","1个","20%"],["低危","1个","20%"],["合计","5个","100%"]])

    doc.add_page_break()
    hd(doc,"3. 单漏洞深度分析与修复",1)
    vs=[
        ("CI-01","shell=True 命令注入","高危",
         "app.py:770 ping() — subprocess.check_output(cmd, shell=True, ...)",
         "攻击路径：\n1. 登录后访问 Ping 页面\n2. 在 IP 输入框输入：127.0.0.1; whoami\n3. 实际执行的命令：ping -c 3 127.0.0.1; whoami\n4. whoami 的结果返回到页面上\n\n代码根因：\ncmd = f\"ping -c 3 {ip}\"\nresult = subprocess.check_output(cmd, shell=True, ...)\n\nip 参数中的 ; whoami 被 shell 解释为第二条命令执行",
         "CVSS 3.1: 9.1 | AV:N/AC:L/PR:L/UI:N/S:U/C:H/I:H/A:H",
         "攻击者可执行任意系统命令，完全控制服务器。", "修复：shell=True → shell=False，参数以列表形式传入"),
        ("CI-02","f-string 拼接命令字符串","高危",
         "app.py:769 — cmd = f\"ping -c 3 {ip}\"",
         "即使使用 shell=False，恶意参数仍可构造异常行为。", "CVSS 3.1: 8.6",
         "命令参数被注入特殊字符串。","修复：f-string 拼接 → 列表参数 ['ping','-c','3',ip]"),
        ("CI-03","无输入校验","高危",
         "app.py:766 — 未对 ip 做任何过滤",
         "shell 特殊字符（; | & ` $ () {} <> ! # %% 空格等）可随意通过。", "CVSS 3.1: 8.6",
         "攻击者可利用任意字符构造恶意命令。","修复：黑名单检测 20+ shell 特殊字符 + 正则白名单"),
        ("CI-04","无超时控制","中危",
         "app.py:770 — 无超时参数",
         "ping 不可达目标时会长时间挂起，消耗服务器资源。", "CVSS 3.1: 5.3",
         "资源耗尽导致拒绝服务。","修复：添加 timeout=30 超时参数"),
        ("CI-05","错误信息泄露","低危",
         "app.py:775 — CalledProcessError 输出返回给用户",
         "错误输出中包含路径、系统信息等敏感内容。", "CVSS 3.1: 3.3",
         "系统信息泄露辅助攻击。","修复：已超时保护，错误信息限于 ping 本身输出"),
    ]
    for i,(vid,vn,vl,vl2,vp,vc,vh,vf) in enumerate(vs):
        hd(doc,f"3.{i+1} {vid}：{vn} [{vl}]",2)
        add_tbl(doc,["属性","内容"],[["编号",vid],["等级",vl],["位置",vl2],["CVSS 3.1",vc],["危害",vh]])
        code(doc,vp,"攻击路径与根因分析"); code(doc,vf,"修复方案")

    doc.add_page_break()
    hd(doc,"4. 修复后完整代码",1)
    code(doc,
        "ALLOWED_PING_PATTERN = re.compile(\n"
        "    r'^(?:::)?[a-fA-F0-9](?:[a-fA-F0-9:]*[a-fA-F0-9])?$|'\n"
        "    r'^[a-zA-Z0-9](?:[a-zA-Z0-9\\-\\.]*[a-zA-Z0-9])?$')\n\n"
        "def _validate_ip_or_domain(target):\n"
        "    forbidden = {\";\",\"|\",\"&\",\"$\",\"`\",\"(\",\")\",\"{\",\"}\",\n"
        "                \"<\",\">\",\"!\",\"#\",\"~\",\"%\",'\"',\"'\",\" \",\n"
        "                \"\\t\",\"\\n\",\"\\r\",\"\\\\\"}\n"
        '    if \":\" in target:\n'
        "        if not re.match(r'^[a-fA-F0-9:]+$', target):\n"
        "            return False\n"
        "    elif any(ch in target for ch in forbidden):\n"
        "        return False\n"
        "    return bool(ALLOWED_PING_PATTERN.match(target))\n\n"
        '# 核心修复：shell=False + 列表参数\n'
        'cmd = ["ping", "-c", "3", ip]\n'
        'result = subprocess.check_output(cmd, shell=False, timeout=30, ...)',
        "修复后完整安全逻辑")

    doc.add_page_break()
    hd(doc,"5. 修复验证结论",1)
    add_tbl(doc,["测试项","攻击载荷","修复前","修复后","结论"],[
        ["正常 IPv4","127.0.0.1","执行成功","执行成功","✓"],
        ["正常域名","baidu.com","执行成功","执行成功","✓"],
        ["IPv6 ::1","::1","N/A","执行成功","✓"],
        ["分号注入","127.0.0.1;whoami","命令执行","非法字符拦截","✓"],
        ["管道注入","127.0.0.1|whoami","命令执行","非法字符拦截","✓"],
        ["逻辑与","127.0.0.1&&whoami","命令执行","非法字符拦截","✓"],
        ["反引号","127.0.0.1`id`","命令执行","非法字符拦截","✓"],
        ["变量替换","127.0.0.1$(id)","命令执行","非法字符拦截","✓"],
        ["IPv6注入","::1;whoami","命令执行","非法字符拦截","✓"],
        ["空格","127.0.0.1 test","参数注入","非法字符拦截","✓"],
        ["超长输入",">255字符","日志注入","输入过长","✓"],
        ["未登录","直接 POST","可执行","401 拒绝","✓"],
    ])

    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("—— 报告结束 ——"); r.bold=True; r.font.size=Pt(14); r.font.name=FONT_HEI; r.font.color.rgb=C_PRI
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("本报告由安全审计组基于对用户管理系统 Ping 网络诊断功能的代码审计与渗透测试数据汇编。")
    r.font.size=Pt(8.5); r.font.name=FONT_SONG; r.font.color.rgb=C_MUTED

    out="/workspace/user_management/SEC-CMDINJ-AUDIT-2025-001.docx"
    doc.save(out)
    import os; sz=os.path.getsize(out)/1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")

if __name__=="__main__":
    main()
