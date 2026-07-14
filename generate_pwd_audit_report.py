#!/usr/bin/env python3
"""密码修改功能 — 漏洞审查与修复报告"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 配色 ──
C_PRI = RGBColor(0x1A, 0x2B, 0x3C)
C_SEC = RGBColor(0x4A, 0x5A, 0x6A)
C_BODY = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x8C, 0x8C, 0x8C)
C_ACC = RGBColor(0x4A, 0x7A, 0xBF)

FONT_HEI = '黑体'
FONT_SONG = '宋体'
FONT_CODE = 'Consolas'

BG_HDR = 'E2E8F0'
BG_ROW = 'F5F7FA'
BG_CODE = 'F0F2F5'


def set_run_font(run, font_name, size=Pt(10), bold=False, color=None):
    run.font.size = size; run.bold = bold
    if color: run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),color)
    c._tc.get_or_add_tcPr().append(s)


def margins(doc, t=2.0, b=2.0, l=2.5, r=2.5):
    for sec in doc.sections:
        sec.top_margin=Cm(t); sec.bottom_margin=Cm(b); sec.left_margin=Cm(l); sec.right_margin=Cm(r)


def hr(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6)
    bdr=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom'); bt.set(qn('w:val'),'single')
    bt.set(qn('w:sz'),'4'); bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),'D0D5DD')
    bdr.append(bt); p._element.get_or_add_pPr().append(bdr)


def add_tbl(doc, headers, rows):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    tp=t._tbl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); t._tbl.insert(0,tp)
    cm=OxmlElement('w:tblCellMar')
    for s in ('top','left','bottom','right'):
        m=OxmlElement(f'w:{s}'); m.set(qn('w:w'),'36'); m.set(qn('w:type'),'dxa'); cm.append(m)
    tp.append(cm)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(h); set_run_font(r,FONT_HEI,Pt(8.5),True,C_PRI); shade(c,BG_HDR)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            r=p.add_run(str(val)); set_run_font(r,FONT_SONG,Pt(8),color=C_BODY)
            if ri%2==1: shade(c,BG_ROW)
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(2); sp.paragraph_format.space_after=Pt(2)


def add_code(doc, code, label=""):
    if label:
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(1); p.paragraph_format.left_indent=Cm(0.3)
        r=p.add_run(label); set_run_font(r,FONT_HEI,Pt(8),True,C_ACC)
    for line in code.rstrip().split('\n'):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.line_spacing=Pt(12)
        r=p.add_run(line); set_run_font(r,FONT_CODE,Pt(7),color=C_BODY)
        sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),BG_CODE)
        p._element.get_or_add_pPr().append(sh)
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(2); sp.paragraph_format.space_after=Pt(2)


def heading(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    sz={1:(18,10,Pt(18)),2:(14,7,Pt(14)),3:(12,5,Pt(12))}.get(level,(10,5,Pt(11)))
    h.paragraph_format.space_before=Pt(sz[0]); h.paragraph_format.space_after=Pt(sz[1])
    clr={1:C_PRI,2:C_SEC,3:C_BODY}.get(level,C_BODY)
    for r in h.runs: set_run_font(r,FONT_HEI,sz[2],True,clr)
    if level==1: hr(doc); return h


def body(doc, text, bold=False, size=10, indent=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); set_run_font(r,FONT_SONG,Pt(size),bold,color=C_BODY)


def set_style(doc):
    s=doc.styles['Normal']; s.font.name=FONT_SONG; s.font.size=Pt(10)
    s.paragraph_format.line_spacing=1.35; s.paragraph_format.space_after=Pt(4); s.font.color.rgb=C_BODY
    rPr=s.element.get_or_add_rPr()
    rFonts=rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts=OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'),FONT_SONG); rFonts.set(qn('w:ascii'),FONT_SONG); rFonts.set(qn('w:hAnsi'),FONT_SONG)


def main():
    doc=Document(); set_style(doc); margins(doc)

    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("密码修改功能\n漏洞审查与修复报告"); r.bold=True; r.font.size=Pt(24); r.font.name=FONT_HEI; r.font.color.rgb=C_PRI
    doc.add_paragraph()
    info=[("文档编号","SEC-PWD-2025-001"),("密级","内部公开"),("版本号","V1.0"),("发布日期","2026-07-14"),("编制人","潘麒宇")]
    tbl=doc.add_table(rows=len(info), cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(info):
        for j,t in enumerate([k,v]):
            c=tbl.rows[i].cells[j]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(t); r.font.size=Pt(11); r.font.name=FONT_SONG
            if j==0: r.bold=True; r.font.color.rgb=C_PRI
            else: r.font.color.rgb=C_ACC
            shade(c,'EDF2F7' if i%2==0 else 'F5F8FB')
    doc.add_page_break()

    heading(doc,"1. 审计概述",1)
    heading(doc,"1.1 项目信息",2)
    add_tbl(doc,["项目","内容"],[
        ["系统名称","用户管理系统"],
        ["审计范围","密码修改功能（/change-password 路由 + profile.html 模板）"],
        ["审计日期","2026 年 7 月 14 日"],
        ["技术栈","Python 3 / Flask 3.x / bcrypt / Werkzeug"],
    ])
    heading(doc,"1.2 漏洞发现方式",2)
    for it in ["代码审计：逐行审查 change-password 路由逻辑，定位 CSRF/Session/表单校验缺失",
               "黑盒渗透：Burp Suite 拦截请求，篡改参数验证越权路径",
               "OWASP 分类：A01:2021 访问控制失效 / A04:2021 不安全设计 / A08:2021 完整性失效"]:
        body(doc,f"  * {it}",size=9.5,indent=0.5)
    heading(doc,"1.3 参考标准",2)
    add_tbl(doc,["标准","适用范围"],[
        ["OWASP Top 10 2021","A01:2021 / A04:2021 / A08:2021"],
        ["CWE-352","跨站请求伪造（CSRF）"],
        ["CWE-639","通过用户控制的键绕过授权"],
    ])

    doc.add_page_break()
    heading(doc,"2. 漏洞发现清单",1)
    heading(doc,"2.1 漏洞总览",2)
    add_tbl(doc,["编号","漏洞名称","等级","所在路由","漏洞类型"],[
        ["PWD-01","CSRF 跨站请求伪造","高危","POST /change-password","CSRF"],
        ["PWD-02","越权修改任意用户密码","高危","POST /change-password","访问控制"],
        ["PWD-03","无原密码验证","高危","POST /change-password","不安全设计"],
        ["PWD-04","无确认密码校验","低危","POST /change-password","输入校验"],
    ])
    heading(doc,"2.2 风险分布",2)
    add_tbl(doc,["风险等级","数量","占比"],[
        ["高危","3个","75%"],
        ["低危","1个","25%"],
        ["合计","4个","100%"],
    ])

    doc.add_page_break()
    heading(doc,"3. 单漏洞深度分析与修复",1)

    vulns=[
        ("PWD-01","CSRF 跨站请求伪造","高危",
         "app.py:607-656 change-password 路由 — 无 CSRF Token 校验",
         "攻击路径：\n1. 攻击者构造恶意 HTML 页面\n2. 页面中嵌入自动提交表单：\n   <form action='http://target/change-password' method='POST'>\n     <input name='username' value='admin'>\n     <input name='new_password' value='hacked123'>\n   </form>\n   <script>document.forms[0].submit()</script>\n3. 已登录 admin 的用户访问该页面\n4. admin 的密码被静默修改为 hacked123",
         "CVSS 3.1: 8.8 | AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:H",
         "攻击者可绕过登录态限制，无需用户交互即可完成密码篡改，结合恶意页面可实现任意账号接管。",
         "表单添加 CSRF Token 隐藏字段：\n<input type='hidden' name='csrf_token' value='{{ csrf_token() }}'>\n服务端校验：\nif not token or token != session.get('_csrf_token'):\n    return render_template(..., error='会话已过期'), 400"),

        ("PWD-02","越权修改任意用户密码","高危",
         "app.py:610 change-password — 直接使用表单 username 不做归属校验",
         "攻击路径：\n1. 登录 alice 账号\n2. 获取个人中心的 CSRF Token\n3. POST 到 /change-password\n4. username=admin&new_password=hacked123\n5. admin 密码被 alice 修改\n\n代码根因：\nusername = request.form.get('username', '').strip()\n# ↑ 直接使用表单提交的username，不与session比对",
         "CVSS 3.1: 7.5 | AV:N/AC:L/PR:L/UI:N/S:U/C:N/I:H/A:N",
         "任何已登录用户可修改系统中任意用户的密码，导致权限完全失控。",
         "添加 username 归属校验：\nform_username = request.form.get('username', '').strip()\nif form_username != current_user:\n    return render_template(..., error='无权修改其他用户的密码'), 403"),

        ("PWD-03","无原密码验证","高危",
         "app.py:610-611 change-password — 未校验原密码",
         "攻击路径：\n1. 获取 CSRF Token\n2. POST 到 /change-password\n3. 直接传入新密码，无需提供原密码\n4. 密码被任意修改\n\n代码根因：\n# 无任何原密码校验逻辑\nUSERS[current_user]['password'] = generate_password_hash(new_password)",
         "CVSS 3.1: 7.5 | AV:N/AC:L/PR:L/UI:N/S:U:C/N:I/H:A:N",
         "任何拿到 CSRF Token 的登录用户可随意修改密码，进一步放大 CSRF 和越权漏洞的危害。",
         "添加原密码验证：\nuser_record = USERS.get(current_user)\nif user_record and not check_password_hash(user_record['password'], old_password):\n    return render_template(..., error='原密码错误')"),

        ("PWD-04","无确认密码校验","低危",
         "app.py:611-612 change-password — 缺少确认密码字段校验",
         "攻击者误输新密码导致无法登录；键盘误触导致密码被改为非预期值。",
         "CVSS 3.1: 3.1 | AV:N/AC:L/PR:L/UI:N/S:U/C:N/I:N/A:L",
         "影响用户体验，可能造成用户无法登录，需管理员介入重置。",
         "添加确认密码字段及校验：\nif new_password != confirm_password:\n    return render_template(..., error='两次输入的密码不一致')"),
    ]

    for idx,(vid,vname,vlevel,vloc,vpath,vcvss,vharm,vfix) in enumerate(vulns):
        heading(doc,f"3.{idx+1} {vid}：{vname} [{vlevel}]",2)
        add_tbl(doc,["属性","内容"],[
            ["漏洞编号",vid],["风险等级",vlevel],["代码位置",vloc],
            ["CVSS 3.1",vcvss],["危害程度",vharm]])
        add_code(doc,vpath,"攻击路径与根因分析")
        add_code(doc,vfix,"修复方案")

    doc.add_page_break()
    heading(doc,"4. 修复后核心代码",1)
    add_code(doc,
        "@app.route(\"/change-password\", methods=[\"POST\"])\n"
        "@login_required\n"
        "def change_password():\n"
        "    current_user = session.get(\"username\")\n"
        "    old_password = request.form.get(\"old_password\", \"\")\n"
        "    new_password = request.form.get(\"new_password\", \"\")\n"
        "    confirm_password = request.form.get(\"confirm_password\", \"\")\n"
        "    token = request.form.get(\"csrf_token\", \"\")\n\n"
        "    # [安全] CSRF 校验\n"
        "    if not token or token != session.get(\"_csrf_token\"):\n"
        "        return \"会话已过期\", 400\n\n"
        "    # [安全] 只能修改自己的密码\n"
        "    form_username = request.form.get(\"username\", \"\").strip()\n"
        "    if form_username != current_user:\n"
        "        return \"无权修改其他用户的密码\", 403\n\n"
        "    # [安全] 验证原密码\n"
        "    user_record = USERS.get(current_user)\n"
        "    if user_record and not check_password_hash(user_record[\"password\"], old_password):\n"
        "        return \"原密码错误\"\n\n"
        "    # [安全] 两次密码一致校验\n"
        "    if new_password != confirm_password:\n"
        "        return \"两次输入的密码不一致\"\n\n"
        "    USERS[current_user][\"password\"] = generate_password_hash(new_password)\n"
        "    return redirect(\"/profile\")",
        "修复后 /change-password 路由")

    doc.add_page_break()
    heading(doc,"5. 修复验证结论",1)
    add_tbl(doc,["测试项","攻击载荷/操作","修复前","修复后","结论"],[
        ["无 CSRF Token","POST 缺 csrf_token","密码被修改","400 拒绝","✓"],
        ["越权改他人密码","alice 改 admin 密码","admin 密码被改","403 拒绝","✓"],
        ["原密码错误","old_password=wrong","密码仍可修改","原密码错误","✓"],
        ["两次密码不一致","new≠confirm","密码修改成功","密码不一致","✓"],
        ["正确修改流程","原密码+新密码+CSRF","N/A","修改成功","✓"],
        ["新密码登录","使用新密码登录","N/A","登录成功","✓"],
    ])

    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("—— 报告结束 ——"); r.bold=True; r.font.size=Pt(14); r.font.name=FONT_HEI; r.font.color.rgb=C_PRI
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("本报告由安全审计组基于对用户管理系统密码修改功能的代码审计与渗透测试数据汇编。")
    r.font.size=Pt(8.5); r.font.name=FONT_SONG; r.font.color.rgb=C_MUTED

    out="/workspace/user_management/SEC-PWD-AUDIT-2025-001.docx"
    doc.save(out)
    import os; sz=os.path.getsize(out)/1024
    print(f"报告已生成：{out} ({sz:.1f} KB)")

if __name__=="__main__":
    main()
